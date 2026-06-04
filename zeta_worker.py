"""Worker module para multiprocessing - debe estar separado del GUI."""
import os
import re
import string

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

def _extract_pdf_text(path):
    if pypdf is None:
        return ""
    text = []
    try:
        with open(path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text.append(t)
    except Exception:
        pass
    return "\n".join(text)

def _extract_docx_text(path):
    if docx is None:
        return ""
    text = []
    try:
        doc = docx.Document(path)
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text.append(cell.text)
    except Exception:
        pass
    return "\n".join(text)

BINARY_EXT = ('.png','.jpg','.jpeg','.gif','.bmp','.ico','.dll','.exe','.so','.dylib',
              '.wav','.ogg','.mp3','.flac','.zip','.rar','.7z','.tar','.gz','.bz2',
              '.mp4','.avi','.mkv','.mov','.pdf','.doc','.docx','.xls','.xlsx',
              '.pyc','.pyo','.class','.o','.obj','.bin','.pack','.idx','.db',
              '.sqlite','.lnk','.pdb','.lib','.exp','.woff','.woff2','.ttf','.otf')

MAX_SIZE = 15 * 1024 * 1024  # 15MB

# Mapear mayúsculas a minúsculas y acentos Latin-1/CP1252 a sus equivalentes ASCII
_FRM = string.ascii_uppercase + "áéíóúñüÁÉÍÓÚÑÜ"
_TO = string.ascii_lowercase + "aeiounuaeiounu"

_TRANS_TABLE = bytes.maketrans(_FRM.encode('latin-1', errors='ignore'), _TO.encode('latin-1', errors='ignore'))

# Secuencias UTF-8 multibyte más comunes en español a sus caracteres ASCII
_UTF8_REPLACEMENTS = [
    (b'\xc3\xa1', b'a'), (b'\xc3\x81', b'a'), # á, Á
    (b'\xc3\xa9', b'e'), (b'\xc3\x89', b'e'), # é, É
    (b'\xc3\xad', b'i'), (b'\xc3\x8d', b'i'), # í, Í
    (b'\xc3\xb3', b'o'), (b'\xc3\x93', b'o'), # ó, Ó
    (b'\xc3\xba', b'u'), (b'\xc3\x9a', b'u'), # ú, Ú
    (b'\xc3\xb1', b'n'), (b'\xc3\x91', b'n'), # ñ, Ñ
    (b'\xc3\xbc', b'u'), (b'\xc3\x9c', b'u'), # ü, Ü
]

def _normalize_bytes(raw: bytes) -> bytes:
    for utf8_seq, ascii_char in _UTF8_REPLACEMENTS:
        raw = raw.replace(utf8_seq, ascii_char)
    return raw.translate(_TRANS_TABLE)

def search_chunk(args):
    """Procesa un lote (chunk) de archivos a nivel de bytes, sin decodificar."""
    paths_meta, keywords_bytes, whole_word, regex, exact = args
    results = []
    new_cache_entries = []
    
    # Precompilar patrones de expresiones regulares si es búsqueda por palabra completa o regex
    if regex:
        try:
            flags = re.IGNORECASE if not exact else 0
            pat = re.compile(keywords_bytes[0], flags)
        except Exception:
            return (len(paths_meta), [], [])
    elif whole_word:
        first_pattern = re.compile(br'(?<![a-z0-9])' + re.escape(keywords_bytes[0]) + br'(?![a-z0-9])')
        other_patterns = [re.compile(br'(?<![a-z0-9])' + re.escape(kw) + br'(?![a-z0-9])') for kw in keywords_bytes[1:]]
    else:
        first_kw = keywords_bytes[0]
        other_kws = keywords_bytes[1:]
        
    for path, mtime, size in paths_meta:
        try:
            if size == 0 or size > MAX_SIZE:
                continue

            ext = os.path.splitext(path)[1].lower()
            if ext == '.pdf' and pypdf is not None:
                text_str = _extract_pdf_text(path)
                raw = text_str.encode('utf-8', errors='ignore')
            elif ext == '.docx' and docx is not None:
                text_str = _extract_docx_text(path)
                raw = text_str.encode('utf-8', errors='ignore')
            else:
                with open(path, 'rb') as f:
                    head = f.read(1024)
                    if head.startswith((b'\xff\xfe', b'\xfe\xff')):
                        pass
                    elif b'\x00' in head:
                        continue
                    rest = f.read()
                    raw = head + rest
                
            normalized = _normalize_bytes(raw)
            normalized_str = normalized.decode('latin-1', errors='ignore')
            new_cache_entries.append((path, mtime, size, normalized_str))
            
            if regex:
                try:
                    count = len(pat.findall(normalized))
                except Exception:
                    count = 0
                if count > 0:
                    results.append((count, path))
            elif whole_word:
                count = len(first_pattern.findall(normalized))
                if count > 0:
                    match_all = True
                    for other_pattern in other_patterns:
                        if not other_pattern.search(normalized):
                            match_all = False
                            break
                    if match_all:
                        results.append((count, path))
            else:
                count = normalized.count(first_kw)
                if count > 0:
                    match_all = True
                    for kw in other_kws:
                        if kw not in normalized:
                            match_all = False
                            break
                    if match_all:
                        results.append((count, path))
        except Exception:
            pass
            
    return (len(paths_meta), results, new_cache_entries)
