import os
import sys
import subprocess
import re
import time
import threading
import queue
import unicodedata
import multiprocessing
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
try:
    from idlelib.tooltip import Hovertip
except ImportError:
    class Hovertip:
        def __init__(self, *args, **kwargs):
            self.text = ""


import ctypes
from ctypes import wintypes
from collections import namedtuple

try:
    version_dll = ctypes.windll.version
    GetFileVersionInfoSizeW = version_dll.GetFileVersionInfoSizeW
    GetFileVersionInfoSizeW.argtypes = [wintypes.LPCWSTR, ctypes.POINTER(wintypes.DWORD)]
    GetFileVersionInfoSizeW.restype = wintypes.DWORD

    GetFileVersionInfoW = version_dll.GetFileVersionInfoW
    GetFileVersionInfoW.argtypes = [wintypes.LPCWSTR, wintypes.DWORD, wintypes.DWORD, wintypes.LPVOID]
    GetFileVersionInfoW.restype = wintypes.BOOL

    VerQueryValueW = version_dll.VerQueryValueW
    VerQueryValueW.argtypes = [wintypes.LPCVOID, wintypes.LPCWSTR, ctypes.POINTER(wintypes.LPVOID), ctypes.POINTER(wintypes.UINT)]
    VerQueryValueW.restype = wintypes.BOOL
except Exception:
    version_dll = None


try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

def _extract_pdf_text(path):
    if pypdf is None:
        return ""
    text = []
    try:
        with open(path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text.append(t)
    except Exception:
        pass
    return "\n".join(text)

def _extract_docx_text(path):
    if docx is None:
        return ""
    text = []
    try:
        doc = docx.Document(path)
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text.append(cell.text)
    except Exception:
        pass
    return "\n".join(text)

# Estructura ligera para simular os.DirEntry en la búsqueda por nombre
FileEntry = namedtuple('FileEntry', ['name', 'path'])

# Estructuras y APIs de Win32 para listar directorios de alto rendimiento
if sys.platform == "win32":
    class WIN32_FIND_DATAW(ctypes.Structure):
        _fields_ = [
            ("dwFileAttributes", wintypes.DWORD),
            ("ftCreationTime", wintypes.FILETIME),
            ("ftLastAccessTime", wintypes.FILETIME),
            ("ftLastWriteTime", wintypes.FILETIME),
            ("nFileSizeHigh", wintypes.DWORD),
            ("nFileSizeLow", wintypes.DWORD),
            ("dwReserved0", wintypes.DWORD),
            ("dwReserved1", wintypes.DWORD),
            ("cFileName", ctypes.c_wchar * 260),
            ("cAlternateFileName", ctypes.c_wchar * 14),
        ]

    class RECT(ctypes.Structure):
        _fields_ = [
            ('left',   ctypes.c_long), ('top',    ctypes.c_long),
            ('right',  ctypes.c_long), ('bottom', ctypes.c_long),
        ]

    class POINT(ctypes.Structure):
        _fields_ = [("x", ctypes.c_long), ("y", ctypes.c_long)]

    try:
        kernel32 = ctypes.windll.kernel32
        user32 = ctypes.windll.user32
        shell32 = ctypes.windll.shell32
        
        try:
            dwmapi = ctypes.windll.dwmapi
        except Exception:
            dwmapi = None

        FindFirstFileW = kernel32.FindFirstFileW
        FindFirstFileW.argtypes = [wintypes.LPCWSTR, ctypes.POINTER(WIN32_FIND_DATAW)]
        FindFirstFileW.restype = wintypes.HANDLE

        FindNextFileW = kernel32.FindNextFileW
        FindNextFileW.argtypes = [wintypes.HANDLE, ctypes.POINTER(WIN32_FIND_DATAW)]
        FindNextFileW.restype = wintypes.BOOL

        FindClose = kernel32.FindClose
        FindClose.argtypes = [wintypes.HANDLE]
        FindClose.restype = wintypes.BOOL

        # user32 signatures
        user32.GetParent.argtypes = [wintypes.HWND]
        user32.GetParent.restype = wintypes.HWND

        user32.GetWindowRect.argtypes = [wintypes.HWND, ctypes.POINTER(RECT)]
        user32.GetWindowRect.restype = wintypes.BOOL

        user32.GetCursorPos.argtypes = [ctypes.POINTER(POINT)]
        user32.GetCursorPos.restype = wintypes.BOOL

        user32.GetAsyncKeyState.argtypes = [ctypes.c_int]
        user32.GetAsyncKeyState.restype = wintypes.SHORT

        user32.ShowWindow.argtypes = [wintypes.HWND, ctypes.c_int]
        user32.ShowWindow.restype = wintypes.BOOL

        user32.SetWindowPos.argtypes = [wintypes.HWND, wintypes.HWND, ctypes.c_int, ctypes.c_int, ctypes.c_int, ctypes.c_int, ctypes.c_uint]
        user32.SetWindowPos.restype = wintypes.BOOL

        user32.IsClipboardFormatAvailable.argtypes = [wintypes.UINT]
        user32.IsClipboardFormatAvailable.restype = wintypes.BOOL

        user32.OpenClipboard.argtypes = [wintypes.HWND]
        user32.OpenClipboard.restype = wintypes.BOOL

        user32.CloseClipboard.argtypes = []
        user32.CloseClipboard.restype = wintypes.BOOL

        user32.GetClipboardData.argtypes = [wintypes.UINT]
        user32.GetClipboardData.restype = wintypes.HANDLE

        # RegisterClassW / CreateWindowExW / SetLayeredWindowAttributes / UnregisterClassW
        user32.RegisterClassW.argtypes = [ctypes.c_void_p]
        user32.RegisterClassW.restype = wintypes.ATOM

        user32.CreateWindowExW.argtypes = [
            wintypes.DWORD, wintypes.LPCWSTR, wintypes.LPCWSTR, wintypes.DWORD,
            ctypes.c_int, ctypes.c_int, ctypes.c_int, ctypes.c_int,
            wintypes.HWND, wintypes.HMENU, wintypes.HINSTANCE, wintypes.LPVOID
        ]
        user32.CreateWindowExW.restype = wintypes.HWND

        user32.SetLayeredWindowAttributes.argtypes = [wintypes.HWND, wintypes.COLORREF, wintypes.BYTE, wintypes.DWORD]
        user32.SetLayeredWindowAttributes.restype = wintypes.BOOL

        user32.UnregisterClassW.argtypes = [wintypes.LPCWSTR, wintypes.HINSTANCE]
        user32.UnregisterClassW.restype = wintypes.BOOL

        # shell32 signatures
        shell32.DragQueryFileW.argtypes = [wintypes.HANDLE, wintypes.UINT, wintypes.LPWSTR, wintypes.UINT]
        shell32.DragQueryFileW.restype = wintypes.UINT

        shell32.DragAcceptFiles.argtypes = [wintypes.HWND, wintypes.BOOL]
        shell32.DragAcceptFiles.restype = None

        shell32.DragFinish.argtypes = [wintypes.HANDLE]
        shell32.DragFinish.restype = None

        # kernel32 signatures
        kernel32.GetModuleHandleW.argtypes = [wintypes.LPCWSTR]
        kernel32.GetModuleHandleW.restype = wintypes.HINSTANCE

        # dwmapi signatures
        if dwmapi:
            dwmapi.DwmSetWindowAttribute.argtypes = [wintypes.HWND, wintypes.DWORD, ctypes.c_void_p, wintypes.DWORD]
            dwmapi.DwmSetWindowAttribute.restype = ctypes.c_long
    except Exception:
        pass

INVALID_HANDLE_VALUE_VAL = ctypes.c_void_p(-1).value if sys.platform == "win32" else -1
FILE_ATTRIBUTE_DIRECTORY = 0x10

# ── Mapeo y traducción de bytes de alto rendimiento para búsqueda sin acentos ─
import string

# Mapear mayúsculas a minúsculas y acentos Latin-1/CP1252 a sus equivalentes ASCII
_FRM = string.ascii_uppercase + "áéíóúñüÁÉÍÓÚÑÜ"
_TO = string.ascii_lowercase + "aeiounuaeiounu"

_TRANS_TABLE = bytes.maketrans(_FRM.encode('latin-1', errors='ignore'), _TO.encode('latin-1', errors='ignore'))

# Secuencias UTF-8 multibyte más comunes en español a sus caracteres ASCII
_UTF8_REPLACEMENTS = [
    (b'\xc3\xa1', b'a'), (b'\xc3\x81', b'a'), # á, Á
    (b'\xc3\xa9', b'e'), (b'\xc3\x89', b'e'), # é, É
    (b'\xc3\xad', b'i'), (b'\xc3\x8d', b'i'), # í, Í
    (b'\xc3\xb3', b'o'), (b'\xc3\x93', b'o'), # ó, Ó
    (b'\xc3\xba', b'u'), (b'\xc3\x9a', b'u'), # ú, Ú
    (b'\xc3\xb1', b'n'), (b'\xc3\x91', b'n'), # ñ, Ñ
    (b'\xc3\xbc', b'u'), (b'\xc3\x9c', b'u'), # ü, Ü
]

def _normalize_bytes(raw: bytes) -> bytes:
    # 1. Reemplazar secuencias multibyte UTF-8 acentuadas
    for utf8_seq, ascii_char in _UTF8_REPLACEMENTS:
        raw = raw.replace(utf8_seq, ascii_char)
    # 2. Traducir mayúsculas y acentos de un solo byte (Latin-1/CP1252) a minúsculas ASCII
    return raw.translate(_TRANS_TABLE)


SKIP_DIRS = {'.git','.godot','node_modules','__pycache__','.vscode','.idea',
             'vnum','$RECYCLE.BIN','System Volume Information','.Trash'}
BLACKLIST_DIRS = ("node_modules", ".git", "appdata", "cache", "$recycle.bin", "system volume information", ".trash")

FILE_TYPES = {
    'type_all': None,
    'type_text': ('.txt', '.log', '.md', '.py', '.js', '.json', '.xml', '.html', '.css', '.cpp', '.c', '.h', '.java', '.docx', '.pdf'),
    'type_images': ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.ico', '.svg'),
    'type_media': ('.mp3', '.wav', '.flac', '.mp4', '.mkv', '.avi', '.mov'),
    'type_sys': ('.exe', '.msi', '.bat', '.cmd', '.dll', '.sys'),
    'type_archive': ('.zip', '.rar', '.7z', '.tar', '.gz')
}

BINARY_EXT = ('.png','.jpg','.jpeg','.gif','.bmp','.ico','.dll','.exe','.so','.dylib',
              '.wav','.ogg','.mp3','.flac','.zip','.rar','.7z','.tar','.gz','.bz2',
              '.mp4','.avi','.mkv','.mov','.pdf','.doc','.docx','.xls','.xlsx',
              '.pyc','.pyo','.class','.o','.obj','.bin','.pack','.idx','.db',
              '.sqlite','.lnk','.pdb','.lib','.exp','.woff','.woff2','.ttf','.otf')
MAX_SIZE = 15 * 1024 * 1024
BATCH_SIZE = 5000  # archivos por bloque enviado al pool
N_WORKERS = max(2, multiprocessing.cpu_count())

TRANSLATIONS = {
    'es': {
        'title': 'Zeta',
        'lbl_base': 'Carpeta Base:',
        'lbl_search': 'Buscar texto:',
        'btn_browse': 'Examinar…',
        'btn_search': 'Buscar',
        'btn_cancel': 'Cancelar',
        'btn_canceling': 'Cancelando…',
        'lbl_options': 'Opciones',
        'opt_type_name': 'Archivo / Carpeta',
        'opt_type_content': 'Solo Contenido',
        'chk_exact': 'Nombre Exacto',
        'chk_whole': 'Palabra Completa',
        'lbl_ignore': 'Ignorar carpetas:',
        'lbl_ignore_custom': 'Otros (comas):',
        'chk_system_trash': 'sistema/papelera',
        'opt_mode_all': 'Todas',
        'opt_mode_any': 'Alguna',
        'lbl_preview': 'Vista Previa',
        'btn_prev': '< Anterior',
        'btn_next': 'Siguiente >',
        'col_name': 'Nombre',
        'col_path': 'Ruta',
        'col_type': 'Tipo',
        'type_folder': 'Carpeta',
        'type_file': 'Archivo',
        'status_ready': 'Listo.',
        'status_searching': 'Buscando…',
        'status_canceled': 'Búsqueda cancelada.',
        'status_no_results': 'No se encontraron coincidencias.',
        'status_results': '✅ {count} resultado(s) en {time:.1f}s',
        'stats_done': 'Total detectados: {total}  │  Leídos: {scanned}  │  Hallados: {matches}  │  Tiempo: {time:.1f}s',
        'stats_running': '⏱ {elapsed:.0f}s  │  Leídos: {scanned} {total_str}  │  Hallados: {matches}',
        'stats_detecting': '{total} detectados…',
        'msg_valid_folder': 'Selecciona una carpeta válida.',
        'msg_valid_query': 'Ingresa texto a buscar.',
        'menu_open': 'Abrir',
        'menu_open_loc': 'Abrir ubicación',
        'menu_copy_path': 'Copiar ruta',
        'status_copied': 'Copiado: {path}',
        'preview_loading': 'Cargando…',
        'preview_folder': '📁 {name}  ({count} elementos)\n\n',
        'preview_more': '\n  … y {count} más.',
        'preview_binary': 'Archivo binario: {name}',
        'preview_truncated': '\n\n… (truncado) …',
        'preview_error': 'Error: {error}',
        'tree_no_results': '❌ Sin resultados',
        'lbl_file_type': 'Tipo de Archivo:',
        'type_all': 'Todo',
        'type_text': 'Texto / Código',
        'type_images': 'Imágenes',
        'type_media': 'Audio / Video',
        'type_sys': 'Ejecutables / Sistema',
        'type_archive': 'Comprimidos',
        'tip_rb_type_name': "Busca coincidencias únicamente en los nombres de archivos y directorios.",
        'tip_rb_type_content': "Busca palabras específicas dentro del texto de los archivos.",
        'tip_chk_exact': "Fuerza a que el nombre coincida de forma idéntica, distinguiendo mayúsculas.",
        'tip_chk_whole': "Evita coincidencias parciales; busca la palabra exacta aislada.",
        'tip_rb_mode_all': "El archivo debe contener todas las palabras ingresadas (Lógica AND).",
        'tip_rb_mode_any': "El archivo puede contener al menos una de las palabras ingresadas (Lógica OR).",
        'tip_cb_file_type': "Filtra los resultados por la categoría de extensión seleccionada.",
        'tip_tree_help': "Tip: Haz clic derecho sobre cualquier archivo de la lista para ver más opciones.",
        'tip_chk_git': "Ignora carpetas de control de versiones (.git).",
        'tip_chk_node': "Ignora carpetas de dependencias de Node.js (node_modules).",
        'tip_chk_appdata': "Ignora la carpeta de datos de aplicaciones de Windows (AppData).",
        'tip_chk_cache': "Ignora carpetas que contengan la palabra 'cache'.",
        'tip_chk_system': "Ignora carpetas de sistema y de reciclaje ($recycle.bin, System Volume Information, .trash).",
        'tip_ignore_custom': "Escribe nombres de carpetas a ignorar, separados por coma (ej: build,dist,temp).",
        'lbl_date': 'Fecha:',
        'date_any': 'Cualquier fecha',
        'date_today': 'Hoy',
        'date_week': 'Esta semana',
        'date_month': 'Este mes',
        'tip_cb_date': 'Filtra los resultados por la fecha de última modificación.',
        'chk_regex': 'Regex',
        'tip_chk_regex': 'Habilita la búsqueda por expresiones regulares.',
        'btn_export': 'Exportar',
        'tip_export': 'Exporta la lista de resultados a un archivo CSV o TXT.',
        'regex_guide_title': 'Guía Rápida de Regex',
        'tip_regex_help': 'Muestra una guía rápida de expresiones regulares.',
        'regex_guide_text': (
            "SÍMBOLOS BÁSICOS:\n"
            "•  . (Punto) -> Coincide con cualquier carácter.\n"
            "•  * (Asterisco) -> Cero o más repeticiones de lo anterior.\n"
            "•  ^ (Sombrerito) -> El texto debe EMPEZAR con esto.\n"
            "•  $ (Signo Pesos) -> El texto debe TERMINAR con esto.\n"
            "•  | (Barra) -> Funciona como un \"O\" lógico.\n\n"
            "EJEMPLOS PRÁCTICOS EN ZETA:\n"
            "1. Filtrar múltiples extensiones a la vez:\n"
            "   \\.(docx|pdf|xlsx)$  -> Busca solo archivos PDF, Word o Excel.\n"
            "2. Archivos que arranquen con un nombre fijo:\n"
            "   ^Zeta.*\\.py$        -> Busca scripts que empiecen con \"Zeta\".\n"
            "3. Buscar respaldos numéricos de 4 dígitos:\n"
            "   backup_\\d{4}        -> Encuentra \"backup_2026\", ignora \"backup_final\".\n"
            "4. Buscar tareas pendientes dentro del código (Solo Contenido):\n"
            "   (?i)TODO:           -> Encuentra marcas \"todo:\" sin importar mayúsculas."
        ),
        'btn_close': 'Cerrar',
        'chk_sound': 'Sonido',
        'tip_chk_sound': 'Emite un sonido cuando finaliza la búsqueda.',
        'status_folder_dropped': 'Carpeta base cambiada a: {path}',
        'menu_add_star': 'Añadir Estrella',
        'menu_remove_star': 'Quitar Estrella',
        'menu_color_tags': 'Etiquetas de Color',
        'menu_remove_tag': 'Quitar Etiqueta',
        'menu_remove_bookmark': 'Eliminar Marcador',
        'tree_no_bookmarks': '❌ Sin marcadores',
        'tab_search': 'Búsqueda',
        'tab_bookmarks': 'Marcadores',
        'tab_space': 'Espacio',
        'btn_analyze_space': 'Analizar Espacio',
        'btn_cancel_space': 'Cancelar',
        'col_size': 'Tamaño',
        'lbl_top_heaviest': 'Top 50 Archivos más Pesados',
        'lbl_categories_breakdown': 'Desglose por Categorías',
        'status_analyzing': 'Analizando espacio (detectados: {count})...',
        'status_analysis_done': '✅ Análisis de espacio completado en {time:.1f}s',
        'cat_documents': 'Documentos',
        'cat_images': 'Imágenes',
        'cat_media': 'Video / Audio',
        'cat_system': 'Ejecutables / Sistema',
        'cat_archive': 'Comprimidos',
        'cat_others': 'Otros'
    },
    'en': {
        'title': 'Zeta',
        'lbl_base': 'Base Folder:',
        'lbl_search': 'Search text:',
        'btn_browse': 'Browse…',
        'btn_search': 'Search',
        'btn_cancel': 'Cancel',
        'btn_canceling': 'Canceling…',
        'lbl_options': 'Options',
        'opt_type_name': 'File / Folder',
        'opt_type_content': 'Content Only',
        'chk_exact': 'Exact Name',
        'chk_whole': 'Whole Word',
        'lbl_ignore': 'Ignore folders:',
        'lbl_ignore_custom': 'Others (commas):',
        'chk_system_trash': 'system/trash',
        'opt_mode_all': 'All',
        'opt_mode_any': 'Any',
        'lbl_preview': 'Preview',
        'btn_prev': '< Previous',
        'btn_next': 'Next >',
        'col_name': 'Name',
        'col_path': 'Path',
        'col_type': 'Type',
        'type_folder': 'Folder',
        'type_file': 'File',
        'status_ready': 'Ready.',
        'status_searching': 'Searching…',
        'status_canceled': 'Search canceled.',
        'status_no_results': 'No matches found.',
        'status_results': '✅ {count} match(es) in {time:.1f}s',
        'stats_done': 'Total detected: {total}  │  Read: {scanned}  │  Found: {matches}  │  Time: {time:.1f}s',
        'stats_running': '⏱ {elapsed:.0f}s  │  Read: {scanned} {total_str}  │  Found: {matches}',
        'stats_detecting': '{total} detected…',
        'msg_valid_folder': 'Select a valid folder.',
        'msg_valid_query': 'Enter text to search.',
        'menu_open': 'Open',
        'menu_open_loc': 'Open location',
        'menu_copy_path': 'Copy path',
        'status_copied': 'Copied: {path}',
        'preview_loading': 'Loading…',
        'preview_folder': '📁 {name}  ({count} items)\n\n',
        'preview_more': '\n  … and {count} more.',
        'preview_binary': 'Binary file: {name}',
        'preview_truncated': '\n\n… (truncated) …',
        'preview_error': 'Error: {error}',
        'tree_no_results': '❌ No results',
        'lbl_file_type': 'File Type:',
        'type_all': 'All',
        'type_text': 'Text / Code',
        'type_images': 'Images',
        'type_media': 'Audio / Video',
        'type_sys': 'Executable / System',
        'type_archive': 'Compressed',
        'tip_rb_type_name': "Searches for matches strictly within file and directory names.",
        'tip_rb_type_content': "Searches for specific keywords inside the text content of files.",
        'tip_chk_exact': "Forces the name to match identically, case-sensitive.",
        'tip_chk_whole': "Prevents partial matches; searches for the exact standalone word.",
        'tip_rb_mode_all': "The file must contain all entered keywords (AND logic).",
        'tip_rb_mode_any': "The file can contain at least one of the entered keywords (OR logic).",
        'tip_cb_file_type': "Filters results by the selected extension category.",
        'tip_tree_help': "Tip: Right-click on any file in the list to view more options.",
        'tip_chk_git': "Ignores version control folders (.git).",
        'tip_chk_node': "Ignores Node.js dependencies folders (node_modules).",
        'tip_chk_appdata': "Ignores Windows application data folder (AppData).",
        'tip_chk_cache': "Ignores folders containing the word 'cache'.",
        'tip_chk_system': "Ignores system and trash folders ($recycle.bin, System Volume Information, .trash).",
        'tip_ignore_custom': "Type folder names to ignore, separated by commas (e.g. build,dist,temp).",
        'lbl_date': 'Date:',
        'date_any': 'Any date',
        'date_today': 'Today',
        'date_week': 'This week',
        'date_month': 'This month',
        'tip_cb_date': 'Filters results by the last modification date.',
        'chk_regex': 'Regex',
        'tip_chk_regex': 'Enables regular expression search.',
        'btn_export': 'Export',
        'tip_export': 'Exports the list of results to a CSV or TXT file.',
        'regex_guide_title': 'Regex Quick Guide',
        'tip_regex_help': 'Shows a quick guide for regular expressions.',
        'regex_guide_text': (
            "BASIC SYMBOLS:\n"
            "•  . (Dot) -> Matches any character.\n"
            "•  * (Asterisk) -> Zero or more repetitions of the preceding element.\n"
            "•  ^ (Caret) -> The text must START with this.\n"
            "•  $ (Dollar sign) -> The text must END with this.\n"
            "•  | (Pipe) -> Works as a logical \"OR\".\n\n"
            "PRACTICAL EXAMPLES IN ZETA:\n"
            "1. Filter multiple extensions at once:\n"
            "   \\.(docx|pdf|xlsx)$  -> Searches only PDF, Word or Excel files.\n"
            "2. Files starting with a specific name:\n"
            "   ^Zeta.*\\.py$        -> Searches for scripts starting with \"Zeta\".\n"
            "3. Search for 4-digit numerical backups:\n"
            "   backup_\\d{4}        -> Finds \"backup_2026\", ignores \"backup_final\".\n"
            "4. Search for pending tasks in code (Content Only):\n"
            "   (?i)TODO:           -> Finds \"todo:\" marks case-insensitively."
        ),
        'btn_close': 'Close',
        'chk_sound': 'Sound',
        'tip_chk_sound': 'Plays a sound when the search completes.',
        'status_folder_dropped': 'Base folder changed to: {path}',
        'menu_add_star': 'Add Star',
        'menu_remove_star': 'Remove Star',
        'menu_color_tags': 'Color Tags',
        'menu_remove_tag': 'Remove Tag',
        'menu_remove_bookmark': 'Remove Bookmark',
        'tree_no_bookmarks': '❌ No bookmarks',
        'tab_search': 'Search',
        'tab_bookmarks': 'Bookmarks',
        'tab_space': 'Space',
        'btn_analyze_space': 'Analyze Space',
        'btn_cancel_space': 'Cancel',
        'col_size': 'Size',
        'lbl_top_heaviest': 'Top 50 Heaviest Files',
        'lbl_categories_breakdown': 'Space by Category',
        'status_analyzing': 'Analyzing space (detected: {count})...',
        'status_analysis_done': '✅ Space analysis completed in {time:.1f}s',
        'cat_documents': 'Documents',
        'cat_images': 'Images',
        'cat_media': 'Video / Audio',
        'cat_system': 'System / Executables',
        'cat_archive': 'Compressed',
        'cat_others': 'Others'
    }
}

# ── helpers ──────────────────────────────────────────────────────────────────
def _normalize(text: str) -> str:
    return unicodedata.normalize('NFKD', text).encode('ASCII', 'ignore').decode().lower()

def _normalize_name(name: str) -> str:
    return re.sub(r'[^a-z0-9]', '', _normalize(name))

def _normalize_regex(pattern: str) -> str:
    def repl(match):
        val = match.group(0)
        if val.startswith('\\'):
            return val
        return _normalize(val)
    return re.sub(r'\\.|.', repl, pattern)

def _format_size(size_bytes):
    if size_bytes < 1024:
        return f"{size_bytes} B"
    elif size_bytes < 1024 * 1024:
        return f"{size_bytes / 1024:.2f} KB"
    elif size_bytes < 1024 * 1024 * 1024:
        return f"{size_bytes / (1024 * 1024):.2f} MB"
    else:
        return f"{size_bytes / (1024 * 1024 * 1024):.2f} GB"


def _is_blacklisted(path: str, blacklist: list) -> bool:
    if not blacklist:
        return False
    parts = os.path.normpath(path).lower().split(os.sep)
    if parts and parts[0].endswith(':'):
        parts = parts[1:]
    for kw in blacklist:
        if '/' in kw or '\\' in kw:
            kw_norm = os.path.normpath(kw).lower()
            path_norm = os.path.normpath(path).lower()
            if kw_norm in path_norm:
                return True
        else:
            if kw in parts:
                return True
    return False

def make_regex_pattern(word: str) -> str:
    char_map = {
        'a':'[aAáÁàÀäÄ]','e':'[eEéÉèÈëË]','i':'[iIíÍìÌïÏ]',
        'o':'[oOóÓòÒöÖ]','u':'[uUúÚùÙüÜ]','n':'[nNñÑ]','c':'[cCçÇ]'
    }
    norm = _normalize(word)
    return ''.join(char_map.get(c, re.escape(c)) for c in norm)

def set_dark_titlebar(window, dark: bool):
    if sys.platform != "win32":
        return
    import ctypes
    try:
        window.update_idletasks()
        hwnd = ctypes.windll.user32.GetParent(window.winfo_id())
        if not hwnd:
            hwnd = window.winfo_id()
        value = ctypes.c_int(1 if dark else 0)
        ctypes.windll.dwmapi.DwmSetWindowAttribute(hwnd, 20, ctypes.byref(value), ctypes.sizeof(value))
        ctypes.windll.dwmapi.DwmSetWindowAttribute(hwnd, 19, ctypes.byref(value), ctypes.sizeof(value))
    except Exception:
        pass

def get_clipboard_files():
    if sys.platform != "win32":
        return []
    CF_HDROP = 15
    user32 = ctypes.windll.user32
    shell32 = ctypes.windll.shell32
    
    try:
        if not user32.IsClipboardFormatAvailable(CF_HDROP):
            return []
            
        if not user32.OpenClipboard(None):
            return []
            
        files = []
        try:
            h_drop = user32.GetClipboardData(CF_HDROP)
            if h_drop:
                num_files = shell32.DragQueryFileW(h_drop, 0xFFFFFFFF, None, 0)
                for i in range(num_files):
                    length = shell32.DragQueryFileW(h_drop, i, None, 0)
                    buf = ctypes.create_unicode_buffer(length + 1)
                    shell32.DragQueryFileW(h_drop, i, buf, length + 1)
                    files.append(buf.value)
        finally:
            user32.CloseClipboard()
        return files
    except Exception:
        return []

# ── Drag & Drop nativo (overlay transparente en hilo Python dedicado) ──────────
# Referencia global para evitar que el garbage collector elimine el hilo.
_dnd_thread = None

def setup_drag_and_drop(window, callback):
    """
    Crea una ventana Win32 overlay transparente en un hilo Python (threading.Thread)
    dedicado, con su propio bucle de mensajes.
    
    Para evitar el congelamiento de la UI por problemas de hit-testing en diferentes hilos
    y permitir la interacción normal con la ventana de Tkinter, el overlay se mantiene
    completamente oculto (SW_HIDE). Se muestra dinámicamente cubriendo la ventana solo
    cuando detectamos que el usuario está arrastrando archivos desde fuera (botón izquierdo
    presionado y cursor dentro de los límites de la ventana sin haber iniciado el click dentro).
    """
    if sys.platform != 'win32':
        return
    global _dnd_thread

    # Obtenemos el HWND nativo de Tkinter en el hilo principal
    window.update_idletasks()
    tk_hwnd = ctypes.windll.user32.GetParent(window.winfo_id())
    if not tk_hwnd:
        tk_hwnd = window.winfo_id()

    _overlay_hwnd = [None]  # contenedor mutable para el HWND del overlay

    # Bandera para indicar si el click empezó en la ventana de Tkinter
    window.tkinter_mouse_down = False

    def _on_press(event):
        window.tkinter_mouse_down = True

    def _on_release(event):
        window.tkinter_mouse_down = False

    window.bind_all("<ButtonPress-1>", _on_press)
    window.bind_all("<ButtonRelease-1>", _on_release)

    def _thread():
        u32 = ctypes.windll.user32
        s32 = ctypes.windll.shell32
        k32 = ctypes.windll.kernel32

        # Firmas explícitas para evitar truncamiento en 64 bits.
        s32.DragAcceptFiles.restype  = None
        s32.DragAcceptFiles.argtypes = [ctypes.c_void_p, ctypes.c_bool]
        s32.DragFinish.restype       = None
        s32.DragFinish.argtypes      = [ctypes.c_void_p]
        s32.DragQueryFileW.restype   = ctypes.c_uint
        s32.DragQueryFileW.argtypes  = [ctypes.c_void_p, ctypes.c_uint,
                                        ctypes.c_wchar_p, ctypes.c_uint]

        WM_DROPFILES      = 0x0233
        WM_DESTROY        = 0x0002
        WS_EX_LAYERED     = 0x00080000
        WS_EX_NOACTIVATE  = 0x08000000
        WS_EX_TOPMOST     = 0x00000008
        WS_POPUP          = 0x80000000
        LWA_ALPHA         = 0x00000002

        WNDPROC = ctypes.WINFUNCTYPE(
            ctypes.c_ssize_t,   # LRESULT
            ctypes.c_void_p,    # HWND
            ctypes.c_uint,      # msg
            ctypes.c_size_t,    # wparam
            ctypes.c_ssize_t,   # lparam
        )

        class WNDCLASSW(ctypes.Structure):
            _fields_ = [
                ('style',          ctypes.c_uint),
                ('lpfnWndProc',    WNDPROC),
                ('cbClsExtra',     ctypes.c_int),
                ('cbWndExtra',     ctypes.c_int),
                ('hInstance',      ctypes.c_void_p),
                ('hIcon',          ctypes.c_void_p),
                ('hCursor',        ctypes.c_void_p),
                ('hbrBackground',  ctypes.c_void_p),
                ('lpszMenuName',   ctypes.c_wchar_p),
                ('lpszClassName',  ctypes.c_wchar_p),
            ]

        class MSG(ctypes.Structure):
            _fields_ = [
                ('hwnd',    ctypes.c_void_p),
                ('message', ctypes.c_uint),
                ('wParam',  ctypes.c_size_t),
                ('lParam',  ctypes.c_ssize_t),
                ('time',    ctypes.c_ulong),
                ('pt',      ctypes.c_long * 2),
            ]

        def wnd_proc(hwnd, msg, wparam, lparam):
            if msg == WM_DROPFILES:
                try:
                    h_drop = wparam
                    num = s32.DragQueryFileW(h_drop, 0xFFFFFFFF, None, 0)
                    path = None
                    if num > 0:
                        length = s32.DragQueryFileW(h_drop, 0, None, 0)
                        buf = ctypes.create_unicode_buffer(length + 1)
                        s32.DragQueryFileW(h_drop, 0, buf, length + 1)
                        path = buf.value
                    s32.DragFinish(h_drop)
                    if path:
                        window.after(0, lambda p=path: callback(p))
                except Exception:
                    pass
                return 0
            elif msg == WM_DESTROY:
                u32.PostQuitMessage(0)
                return 0
            return u32.DefWindowProcW(hwnd, msg, wparam, lparam)

        u32.DefWindowProcW.restype  = ctypes.c_ssize_t
        u32.DefWindowProcW.argtypes = [
            ctypes.c_void_p,
            ctypes.c_uint,
            ctypes.c_size_t,
            ctypes.c_ssize_t,
        ]

        wnd_proc_ref  = WNDPROC(wnd_proc)
        h_instance    = k32.GetModuleHandleW(None)
        class_name    = f"ZetaDnD{abs(id(window))}"

        wc = WNDCLASSW()
        wc.lpfnWndProc   = wnd_proc_ref
        wc.hInstance     = h_instance
        wc.lpszClassName = class_name
        u32.RegisterClassW(ctypes.byref(wc))

        # Creamos la ventana inicialmente oculta y sin tamaño
        hwnd = u32.CreateWindowExW(
            WS_EX_LAYERED | WS_EX_NOACTIVATE | WS_EX_TOPMOST,
            class_name, None,
            WS_POPUP,
            0, 0, 1, 1,
            None, None, h_instance, None
        )
        _overlay_hwnd[0] = hwnd

        u32.SetLayeredWindowAttributes(hwnd, 0, 1, LWA_ALPHA)
        u32.ShowWindow(hwnd, 0)  # SW_HIDE (inicialmente oculta)
        s32.DragAcceptFiles(hwnd, True)

        msg_obj = MSG()
        while True:
            ret = u32.GetMessageW(ctypes.byref(msg_obj), None, 0, 0)
            if ret <= 0:
                break
            u32.TranslateMessage(ctypes.byref(msg_obj))
            u32.DispatchMessageW(ctypes.byref(msg_obj))

        u32.UnregisterClassW(class_name, h_instance)

    class POINT(ctypes.Structure):
        _fields_ = [("x", ctypes.c_long), ("y", ctypes.c_long)]

    class RECT(ctypes.Structure):
        _fields_ = [
            ('left',   ctypes.c_long), ('top',    ctypes.c_long),
            ('right',  ctypes.c_long), ('bottom', ctypes.c_long),
        ]

    def check_drag():
        hwnd = _overlay_hwnd[0]
        if not hwnd:
            window.after(100, check_drag)
            return
        try:
            # 1. Comprobar si el botón izquierdo del ratón está físicamente presionado
            is_lbutton_down = (ctypes.windll.user32.GetAsyncKeyState(0x01) & 0x8000) != 0

            # 2. Comprobar si el cursor está dentro del área de la ventana principal
            cursor_pos = POINT()
            ctypes.windll.user32.GetCursorPos(ctypes.byref(cursor_pos))

            rect = RECT()
            ctypes.windll.user32.GetWindowRect(ctypes.c_void_p(tk_hwnd), ctypes.byref(rect))

            is_over = (rect.left <= cursor_pos.x <= rect.right and
                       rect.top <= cursor_pos.y <= rect.bottom)

            # Si el mouse está abajo, el cursor está dentro y el click no se inició en Tkinter,
            # significa que están arrastrando algo desde fuera. Mostramos el overlay.
            if is_lbutton_down and is_over and not window.tkinter_mouse_down:
                ctypes.windll.user32.ShowWindow(hwnd, 4)  # SW_SHOWNOACTIVATE
                ctypes.windll.user32.SetWindowPos(
                    hwnd, ctypes.c_void_p(-1),  # HWND_TOPMOST
                    rect.left, rect.top,
                    rect.right - rect.left,
                    rect.bottom - rect.top,
                    0x0010  # SWP_NOACTIVATE
                )
            else:
                # Si no, mantenemos el overlay oculto para no entorpecer la UI de Tkinter
                ctypes.windll.user32.ShowWindow(hwnd, 0)  # SW_HIDE
        except Exception:
            pass
        window.after(100, check_drag)

    window.after(100, check_drag)

    _dnd_thread = threading.Thread(target=_thread, daemon=True)
    _dnd_thread.start()

# ── pool-worker (debe ser pickleable = nivel de módulo) ───────────────────────
def _search_chunk(args):
    """Procesa un lote (chunk) de archivos a nivel de bytes, sin decodificar."""
    paths_meta, keywords_bytes, whole_word, regex, exact = args
    results = []
    new_cache_entries = []
    
    # Precompilar patrones de expresiones regulares si es búsqueda por palabra completa
    if regex:
        try:
            flags = re.IGNORECASE if not exact else 0
            pat = re.compile(keywords_bytes[0], flags)
        except Exception:
            return (len(paths_meta), [], [])
    elif whole_word:
        first_pattern = re.compile(br'(?<![a-z0-9])' + re.escape(keywords_bytes[0]) + br'(?![a-z0-9])')
        other_patterns = [re.compile(br'(?<![a-z0-9])' + re.escape(kw) + br'(?![a-z0-9])') for kw in keywords_bytes[1:]]
    else:
        first_kw = keywords_bytes[0]
        other_kws = keywords_bytes[1:]
        
    for path, mtime, size in paths_meta:
        try:
            if size == 0 or size > MAX_SIZE:
                continue

            ext = os.path.splitext(path)[1].lower()
            if ext == '.pdf' and pypdf is not None:
                text_str = _extract_pdf_text(path)
                raw = text_str.encode('utf-8', errors='ignore')
            elif ext == '.docx' and docx is not None:
                text_str = _extract_docx_text(path)
                raw = text_str.encode('utf-8', errors='ignore')
            else:
                with open(path, 'rb') as f:
                    head = f.read(1024)
                    if head.startswith((b'\xff\xfe', b'\xfe\xff')):
                        pass
                    elif b'\x00' in head:
                        continue
                    rest = f.read()
                    raw = head + rest
                
            normalized = _normalize_bytes(raw)
            normalized_str = normalized.decode('latin-1', errors='ignore')
            new_cache_entries.append((path, mtime, size, normalized_str))
            
            if regex:
                try:
                    count = len(pat.findall(normalized))
                except Exception:
                    count = 0
                if count > 0:
                    results.append((count, path))
            elif whole_word:
                count = len(first_pattern.findall(normalized))
                if count > 0:
                    match_all = True
                    for other_pattern in other_patterns:
                        if not other_pattern.search(normalized):
                            match_all = False
                            break
                    if match_all:
                        results.append((count, path))
            else:
                count = normalized.count(first_kw)
                if count > 0:
                    match_all = True
                    for kw in other_kws:
                        if kw not in normalized:
                            match_all = False
                            break
                    if match_all:
                        results.append((count, path))
        except Exception:
            pass
            
    return (len(paths_meta), results, new_cache_entries)

# ── lógica de búsqueda ────────────────────────────────────────────────────────
class SearchEngine:
    def __init__(self):
        self.cancel = False
        self._pool = None
        self.config_dir = os.path.join(os.path.expanduser('~'), '.zeta_search')
        self.db_path = os.path.join(self.config_dir, 'zeta_cache.db')
        self._init_db()

    def _init_db(self):
        import sqlite3
        try:
            os.makedirs(self.config_dir, exist_ok=True)
            conn = sqlite3.connect(self.db_path)
            c = conn.cursor()
            c.execute('''
                CREATE TABLE IF NOT EXISTS cache (
                    path TEXT PRIMARY KEY,
                    mtime REAL,
                    size INTEGER,
                    normalized_text TEXT
                )
            ''')
            conn.commit()
            conn.close()
        except Exception:
            pass

    def stop(self):
        self.cancel = True
        if self._pool:
            self._pool.terminate()
            self._pool = None

    def _walk_files_win32(self, root, text_only=True, extensions=None, blacklist=None):
        stack = [root]
        fd = WIN32_FIND_DATAW()
        invalid_val = INVALID_HANDLE_VALUE_VAL
        
        while stack:
            current = stack.pop()
            if self.cancel:
                return
            
            # Folder Blacklist check on the popped directory absolute path
            if _is_blacklisted(current, blacklist):
                continue
                
            search_path = os.path.join(current, "*")
            h = FindFirstFileW(search_path, ctypes.byref(fd))
            if h is None or h == invalid_val:
                continue
                
            try:
                while True:
                    if self.cancel:
                        return
                    name = fd.cFileName
                    if name not in (".", ".."):
                        attrs = fd.dwFileAttributes
                        is_dir = bool(attrs & FILE_ATTRIBUTE_DIRECTORY)
                        full_path = os.path.join(current, name)
                        
                        if is_dir:
                            if name.lower() not in SKIP_DIRS:
                                # Folder Blacklist check before appending to stack
                                if not _is_blacklisted(full_path, blacklist):
                                    stack.append(full_path)
                        else:
                            # Blacklist check for files
                            if _is_blacklisted(full_path, blacklist):
                                if not FindNextFileW(h, ctypes.byref(fd)):
                                    break
                                continue

                            file_size = (fd.nFileSizeHigh * 4294967296) + fd.nFileSizeLow
                            # 15 MB native file size check for content search only
                            if text_only and file_size > MAX_SIZE:
                                if not FindNextFileW(h, ctypes.byref(fd)):
                                    break
                                continue
                            
                            # Extension filter check
                            if extensions is not None:
                                if not name.lower().endswith(extensions):
                                    if not FindNextFileW(h, ctypes.byref(fd)):
                                        break
                                    continue
                                    
                            if text_only:
                                ft = (fd.ftLastWriteTime.dwHighDateTime << 32) + fd.ftLastWriteTime.dwLowDateTime
                                mtime = (ft / 10000000.0) - 11644473600.0
                                yield (full_path, mtime, file_size)
                            else:
                                yield FileEntry(name, full_path)
                                
                    if not FindNextFileW(h, ctypes.byref(fd)):
                        break
            finally:
                FindClose(h)
                                    
    def _walk_files_scandir(self, root, text_only=True, extensions=None, blacklist=None):
        stack = [root]
        while stack:
            current = stack.pop()
            if self.cancel:
                return
            
            # Folder Blacklist check on the popped directory absolute path
            if _is_blacklisted(current, blacklist):
                continue
                
            try:
                with os.scandir(current) as it:
                    for e in it:
                        if self.cancel:
                            return
                        if e.is_dir(follow_symlinks=False):
                            if e.name.lower() not in SKIP_DIRS:
                                # Folder Blacklist check before appending to stack
                                if not _is_blacklisted(e.path, blacklist):
                                    stack.append(e.path)
                        else:
                            # Blacklist check for files
                            if _is_blacklisted(e.path, blacklist):
                                continue

                            try:
                                stat = e.stat(follow_symlinks=False)
                                file_size = stat.st_size
                                mtime = stat.st_mtime
                            except Exception:
                                continue
                                
                            # 15 MB file size check for content search only
                            if text_only and file_size > MAX_SIZE:
                                continue
                                
                            # Extension filter check
                            if extensions is not None:
                                if not e.name.lower().endswith(extensions):
                                    continue
                                    
                            if text_only:
                                yield (e.path, mtime, file_size)
                            else:
                                yield e
            except PermissionError:
                pass

    # ── recorrer disco iterativo (sin recursión) ──────────────────────────────
    def _walk_files(self, root, text_only=True, extensions=None, blacklist=None):
        """Generador iterativo con Win32 API en Windows y fallback a os.scandir."""
        if sys.platform == "win32":
            try:
                yield from self._walk_files_win32(root, text_only, extensions, blacklist)
                return
            except Exception:
                pass
        yield from self._walk_files_scandir(root, text_only, extensions, blacklist)

    def _walk_space_win32(self, root, blacklist=None):
        stack = [root]
        fd = WIN32_FIND_DATAW()
        invalid_val = INVALID_HANDLE_VALUE_VAL
        
        while stack:
            current = stack.pop()
            if self.cancel:
                return
            
            if _is_blacklisted(current, blacklist):
                continue
                
            search_path = os.path.join(current, "*")
            h = FindFirstFileW(search_path, ctypes.byref(fd))
            if h is None or h == invalid_val:
                continue
                
            try:
                while True:
                    if self.cancel:
                        return
                    name = fd.cFileName
                    if name not in (".", ".."):
                        attrs = fd.dwFileAttributes
                        is_dir = bool(attrs & FILE_ATTRIBUTE_DIRECTORY)
                        full_path = os.path.join(current, name)
                        
                        if is_dir:
                            if name.lower() not in SKIP_DIRS:
                                if not _is_blacklisted(full_path, blacklist):
                                    stack.append(full_path)
                        else:
                            if _is_blacklisted(full_path, blacklist):
                                if not FindNextFileW(h, ctypes.byref(fd)):
                                    break
                                continue

                            file_size = (fd.nFileSizeHigh * 4294967296) + fd.nFileSizeLow
                            yield (name, full_path, file_size)
                                
                    if not FindNextFileW(h, ctypes.byref(fd)):
                        break
            finally:
                FindClose(h)

    def _walk_space_scandir(self, root, blacklist=None):
        stack = [root]
        while stack:
            current = stack.pop()
            if self.cancel:
                return
            
            if _is_blacklisted(current, blacklist):
                continue
                
            try:
                with os.scandir(current) as it:
                    for e in it:
                        if self.cancel:
                            return
                        if e.is_dir(follow_symlinks=False):
                            if e.name.lower() not in SKIP_DIRS:
                                if not _is_blacklisted(e.path, blacklist):
                                    stack.append(e.path)
                        else:
                            if _is_blacklisted(e.path, blacklist):
                                continue

                            try:
                                file_size = e.stat(follow_symlinks=False).st_size
                                yield (e.name, e.path, file_size)
                            except Exception:
                                continue
            except PermissionError:
                pass

    def walk_space(self, root, blacklist=None):
        if sys.platform == "win32":
            try:
                yield from self._walk_space_win32(root, blacklist)
                return
            except Exception:
                pass
        yield from self._walk_space_scandir(root, blacklist)


    # ── búsqueda por nombre ───────────────────────────────────────────────────
    def search_name(self, root, keywords, exact, extensions, match_cb, scan_cb, blacklist=None, threshold=None, regex=False):
        """scan_cb(path, scanned, total, matches) - total=-1 mientras recolecta"""
        self.cancel = False
        if regex:
            try:
                flags = 0 if exact else re.IGNORECASE
                pat = re.compile(keywords[0], flags)
            except Exception:
                scan_cb('', 0, 0, 0)
                return
        else:
            kws = [kw.lower() for kw in keywords] if exact else [_normalize_name(kw) for kw in keywords]
        n = 0
        matches = 0
        for entry in self._walk_files(root, text_only=False, extensions=extensions, blacklist=blacklist):
            if self.cancel:
                break
            n += 1
            if regex:
                try:
                    matched = bool(pat.search(entry.name))
                except Exception:
                    matched = False
            else:
                name = entry.name.lower() if exact else _normalize_name(entry.name)
                matched = all(k in name for k in kws)

            if matched:
                if threshold is not None:
                    try:
                        mtime = os.path.getmtime(entry.path)
                        if mtime < threshold:
                            continue
                    except Exception:
                        continue
                score = 100 if (not regex and entry.name.lower() == keywords[0].lower()) else 10
                match_cb((score, entry.path))
                matches += 1
            if n % 2000 == 0:
                scan_cb(entry.path, n, -1, matches)
        scan_cb('', n, n, matches)

    # ── búsqueda por contenido con multiprocessing ────────────────────────────
    def search_content(self, root, keywords, mode_and, whole_word, extensions, match_cb, scan_cb, blacklist=None, threshold=None, regex=False, exact=False):
        """
        Fase 1: recolectar TODOS los archivos (rápido, Win32 / os.scandir).
        Fase 2: procesar con pool de forma dinámica a nivel de bytes puros.
        scan_cb(path, scanned, total, matches)
        """
        self.cancel = False
        kws_norm = [_normalize_regex(kw) for kw in keywords] if regex else [_normalize(kw) for kw in keywords]

        # ── FASE 1: recolectar todos los archivos ──────────────────────────────
        all_files = []
        for path, mtime, size in self._walk_files(root, text_only=True, extensions=extensions, blacklist=blacklist):
            if self.cancel:
                return
            if threshold is not None and mtime < threshold:
                continue
            all_files.append((path, mtime, size))
            # Actualizar UI durante recolección cada 5000 archivos
            if len(all_files) % 5000 == 0:
                scan_cb(path, 0, len(all_files), 0)

        total = len(all_files)
        if total == 0 or self.cancel:
            scan_cb('', 0, 0, 0)
            return

        # Emitir total real inmediatamente
        scan_cb(all_files[-1][0], 0, total, 0)

        # ── FASE 1.5: Cargar caché de SQLite para coincidencia rápida ──────────
        import sqlite3
        cache_map = {}
        try:
            conn = sqlite3.connect(self.db_path)
            c = conn.cursor()
            c.execute("SELECT path, mtime, size, normalized_text FROM cache WHERE path >= ? AND path < ?", (root, root + '\U0010FFFF'))
            for row in c.fetchall():
                cache_map[row[0]] = (row[1], row[2], row[3])
            conn.close()
        except Exception:
            pass

        # ── FASE 2: Validar caché y dividir en cache-hits e uncached_files ────
        uncached_files = []
        done = 0
        matches = 0
        last_path = all_files[-1][0]

        # Precompilar patrones de expresiones regulares para coincidencia de caché
        if regex:
            try:
                query_norm = _normalize_regex(keywords[0])
                flags = re.IGNORECASE if not exact else 0
                pat_str = re.compile(query_norm, flags)
            except Exception:
                scan_cb('', 0, 0, 0)
                return
        elif whole_word:
            first_pattern = re.compile(r'(?<![a-z0-9])' + re.escape(kws_norm[0]) + r'(?![a-z0-9])')
            other_patterns = [re.compile(r'(?<![a-z0-9])' + re.escape(kw) + r'(?![a-z0-9])') for kw in kws_norm[1:]]
        else:
            first_kw = kws_norm[0]
            other_kws = kws_norm[1:]

        for path, mtime, size in all_files:
            if self.cancel:
                return

            cached = cache_map.get(path)
            if cached and cached[0] == mtime and cached[1] == size:
                done += 1
                normalized_str = cached[2]
                if regex:
                    try:
                        count = len(pat_str.findall(normalized_str))
                    except Exception:
                        count = 0
                    if count > 0:
                        matches += 1
                        match_cb((count, path))
                elif whole_word:
                    count = len(first_pattern.findall(normalized_str))
                    if count > 0:
                        match_all = True
                        for other_pattern in other_patterns:
                            if not other_pattern.search(normalized_str):
                                match_all = False
                                break
                        if match_all:
                            matches += 1
                            match_cb((count, path))
                else:
                    count = normalized_str.count(first_kw)
                    if count > 0:
                        match_all = True
                        for kw in other_kws:
                            if kw not in normalized_str:
                                match_all = False
                                break
                        if match_all:
                            matches += 1
                            match_cb((count, path))
            else:
                uncached_files.append((path, mtime, size))

        # Reportar progreso tras evaluar la caché
        if done > 0:
            scan_cb(last_path, done, total, matches)

        # Si no hay archivos sin caché, terminar
        if not uncached_files:
            scan_cb('', done, total, matches)
            return

        # ── FASE 3: procesar archivos no en caché mediante pool de procesos ────
        keywords_bytes = [kw.encode('latin-1', errors='ignore') for kw in kws_norm]
        chunks = [uncached_files[i:i + BATCH_SIZE] for i in range(0, len(uncached_files), BATCH_SIZE)]
        tasks_args = [
            (chunk, keywords_bytes, whole_word, regex, exact)
            for chunk in chunks
        ]

        self._pool = multiprocessing.Pool(processes=N_WORKERS)
        pool = self._pool

        try:
            # chunksize=1 fuerza a que cada proceso trabaje de forma dinámica a demanda con su lote de archivos
            for num_scanned, match_list, new_cache_entries in pool.imap_unordered(_search_chunk, tasks_args, chunksize=1):
                if self.cancel:
                    break
                done += num_scanned
                for match in match_list:
                    matches += 1
                    match_cb(match)
                
                # Cargar nuevas entradas a la caché SQLite
                if new_cache_entries:
                    try:
                        conn = sqlite3.connect(self.db_path)
                        c = conn.cursor()
                        c.executemany("INSERT OR REPLACE INTO cache (path, mtime, size, normalized_text) VALUES (?, ?, ?, ?)", new_cache_entries)
                        conn.commit()
                        conn.close()
                    except Exception:
                        pass

                # Reportar progreso periódicamente tras procesar cada lote
                scan_cb(last_path, done, total, matches)
        finally:
            pool.close()
            pool.join()
            self._pool = None
            if not self.cancel:
                scan_cb('', done, total, matches)


# ── GUI ───────────────────────────────────────────────────────────────────────
class ZetaApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Zeta")
        self.root.geometry("1100x650")
        self.root.minsize(900, 500)

        if sys.platform == "win32":
            import ctypes
            try: ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID('zeta.search.app')
            except: pass

        base = sys._MEIPASS if getattr(sys, 'frozen', False) else os.path.dirname(os.path.abspath(__file__))
        ico = os.path.join(base, "buscador.ico")
        if os.path.exists(ico):
            self.root.iconbitmap(ico)

        self.engine = SearchEngine()
        self.q = queue.Queue()
        self.is_searching = False
        self.last_query = ""
        self.last_type = "nombre"
        self.start_time = 0
        self.stats = {'scanned': 0, 'matches': 0, 'last_path': ''}
        self.is_analyzing_space = False
        self.space_stats = {'scanned': 0, 'total_size': 0}
        self.top50_data = []
        self.category_sizes = {
            'documents': 0, 'images': 0, 'media': 0, 'system': 0, 'archive': 0, 'others': 0
        }
        self.category_colors = {
            'documents': '#f59e0b',
            'images': '#10b981',
            'media': '#3b82f6',
            'system': '#ef4444',
            'archive': '#8b5cf6',
            'others': '#6b7280'
        }
        self.match_indices = []
        self.cur_match = 0
        self.matches_shown = 0

        self.language = "es"
        self.theme = "light"
        self._load_settings()
        self.var_sound = tk.BooleanVar(value=self.sound)
        self.current_file_type_key = 'type_all'
        self.current_date_filter_key = 'date_any'

        # Load custom theme image
        base = sys._MEIPASS if getattr(sys, 'frozen', False) else os.path.dirname(os.path.abspath(__file__))
        self.theme_img_path = os.path.join(base, "theme.png")
        self.theme_photo_light = None
        self.theme_photo_dark = None
        self._load_theme_images()

        self._build_ui()
        self._refresh_bookmarks_tree()
        self.root.after(80, self._poll_queue)
        # Drag & Drop: activar después de que la UI esté construida.
        self.root.after(500, lambda: setup_drag_and_drop(self.root, self._on_folder_drop))

    # ── UI ────────────────────────────────────────────────────────────────────
    def _build_ui(self):
        self.style = ttk.Style()
        self.style.theme_use('clam')

        top = ttk.Frame(self.root, padding=10)
        top.pack(fill=tk.X)

        self.lbl_base_folder = ttk.Label(top, text="Carpeta Base:", font=("Segoe UI", 10, "bold"))
        self.lbl_base_folder.grid(row=0, column=0, sticky=tk.W)
        self.var_path = tk.StringVar()
        self.ent_path = ttk.Entry(top, textvariable=self.var_path, font=("Segoe UI", 10))
        self.ent_path.grid(row=0, column=1, sticky=tk.EW, padx=8)
        self.ent_path.bind("<Control-v>", self._on_paste)
        self.ent_path.bind("<Control-V>", self._on_paste)
        self.btn_browse = ttk.Button(top, text="Examinar…", command=self._browse)
        self.btn_browse.grid(row=0, column=2)
        top.columnconfigure(1, weight=1)

        self.lbl_search_text = ttk.Label(top, text="Buscar texto:", font=("Segoe UI", 10, "bold"))
        self.lbl_search_text.grid(row=1, column=0, sticky=tk.W, pady=6)
        self.var_q = tk.StringVar()
        self.cb_q = ttk.Combobox(top, textvariable=self.var_q, font=("Segoe UI", 10))
        self.cb_q.grid(row=1, column=1, sticky=tk.EW, padx=8, pady=6)
        self.cb_q.bind("<Return>", lambda e: self._start())
        self.cb_q.config(values=self.search_history)
        self.btn_search = ttk.Button(top, text="Buscar", command=self._start)
        self.btn_search.grid(row=1, column=2)

        self.lbl_opts = ttk.LabelFrame(top, text="Opciones", padding=8)
        self.lbl_opts.grid(row=2, column=0, columnspan=3, sticky=tk.EW, pady=8)

        row1 = ttk.Frame(self.lbl_opts)
        row1.pack(fill=tk.X, expand=True, pady=2)

        row2 = ttk.Frame(self.lbl_opts)
        row2.pack(fill=tk.X, expand=True, pady=2)

        row3 = ttk.Frame(self.lbl_opts)
        row3.pack(fill=tk.X, expand=True, pady=2)

        self.var_type = tk.StringVar(value="nombre")
        self.rb_type_name = ttk.Radiobutton(row1, text="Archivo / Carpeta", variable=self.var_type, value="nombre")
        self.rb_type_name.pack(side=tk.LEFT, padx=5)
        self.rb_type_content = ttk.Radiobutton(row1, text="Solo Contenido", variable=self.var_type, value="contenido")
        self.rb_type_content.pack(side=tk.LEFT, padx=5)

        ttk.Separator(row1, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=8)

        self.var_exact = tk.BooleanVar()
        self.chk_exact = ttk.Checkbutton(row1, text="Nombre Exacto", variable=self.var_exact)
        self.chk_exact.pack(side=tk.LEFT, padx=5)
        self.var_whole = tk.BooleanVar()
        self.chk_whole = ttk.Checkbutton(row1, text="Palabra Completa", variable=self.var_whole)
        self.chk_whole.pack(side=tk.LEFT, padx=5)
        self.var_regex = tk.BooleanVar()
        self.chk_regex = ttk.Checkbutton(row1, text="Regex", variable=self.var_regex)
        self.chk_regex.pack(side=tk.LEFT, padx=5)
        self.lbl_regex_help = ttk.Label(row1, text="ℹ️", cursor="hand2")
        self.lbl_regex_help.pack(side=tk.LEFT, padx=(0, 5))
        self.lbl_regex_help.bind("<Button-1>", lambda e: self._show_regex_guide())

        ttk.Separator(row1, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=8)

        self.var_mode = tk.StringVar(value="and")
        self.rb_mode_all = ttk.Radiobutton(row1, text="Todas", variable=self.var_mode, value="and")
        self.rb_mode_all.pack(side=tk.LEFT, padx=4)
        self.rb_mode_any = ttk.Radiobutton(row1, text="Alguna", variable=self.var_mode, value="or")
        self.rb_mode_any.pack(side=tk.LEFT, padx=4)

        # Row 2: Filters and Preferences
        self.lbl_file_type = ttk.Label(row2, text="Tipo:")
        self.lbl_file_type.pack(side=tk.LEFT, padx=4)
        self.var_file_type = tk.StringVar()
        self.cb_file_type = ttk.Combobox(row2, textvariable=self.var_file_type, state="readonly", width=12)
        self.cb_file_type.pack(side=tk.LEFT, padx=4)
        self.cb_file_type.bind("<<ComboboxSelected>>", lambda e: self._on_file_type_change())

        self.lbl_date = ttk.Label(row2, text="Fecha:")
        self.lbl_date.pack(side=tk.LEFT, padx=4)
        self.var_date = tk.StringVar()
        self.cb_date = ttk.Combobox(row2, textvariable=self.var_date, state="readonly", width=14)
        self.cb_date.pack(side=tk.LEFT, padx=4)
        self.cb_date.bind("<<ComboboxSelected>>", lambda e: self._on_date_change())

        # Theme button (small, modern)
        self.btn_theme = ttk.Button(row2, command=self._toggle_theme)
        self.btn_theme.pack(side=tk.RIGHT, padx=5)
        
        # Language Selector Combobox
        lang_display = "Español" if self.language == "es" else "English"
        self.var_lang = tk.StringVar(value=lang_display)
        self.cb_lang = ttk.Combobox(row2, textvariable=self.var_lang, values=["Español", "English"], width=8, state="readonly")
        self.cb_lang.pack(side=tk.RIGHT, padx=5)
        self.cb_lang.bind("<<ComboboxSelected>>", lambda e: self._on_language_change())

        # --- Row 3: Blacklist Exclusions ---
        self.lbl_ignore = ttk.Label(row3, text="Ignorar:", font=("Segoe UI", 9, "bold"))
        self.lbl_ignore.pack(side=tk.LEFT, padx=(5, 8))

        self.var_ignore_git = tk.BooleanVar(value=self.ignore_git)
        self.chk_ignore_git = ttk.Checkbutton(row3, text=".git", variable=self.var_ignore_git)
        self.chk_ignore_git.pack(side=tk.LEFT, padx=5)

        self.var_ignore_node = tk.BooleanVar(value=self.ignore_node)
        self.chk_ignore_node = ttk.Checkbutton(row3, text="node_modules", variable=self.var_ignore_node)
        self.chk_ignore_node.pack(side=tk.LEFT, padx=5)

        self.var_ignore_appdata = tk.BooleanVar(value=self.ignore_appdata)
        self.chk_ignore_appdata = ttk.Checkbutton(row3, text="AppData", variable=self.var_ignore_appdata)
        self.chk_ignore_appdata.pack(side=tk.LEFT, padx=5)

        self.var_ignore_cache = tk.BooleanVar(value=self.ignore_cache)
        self.chk_ignore_cache = ttk.Checkbutton(row3, text="cache", variable=self.var_ignore_cache)
        self.chk_ignore_cache.pack(side=tk.LEFT, padx=5)

        self.var_ignore_system = tk.BooleanVar(value=self.ignore_system)
        self.chk_ignore_system = ttk.Checkbutton(row3, text="sistema/papelera", variable=self.var_ignore_system)
        self.chk_ignore_system.pack(side=tk.LEFT, padx=5)

        ttk.Separator(row3, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=8)

        self.lbl_ignore_custom = ttk.Label(row3, text="Otros (comas):")
        self.lbl_ignore_custom.pack(side=tk.LEFT, padx=4)

        self.var_custom_ignore = tk.StringVar(value=self.custom_ignore)
        self.ent_custom_ignore = ttk.Entry(row3, textvariable=self.var_custom_ignore, font=("Segoe UI", 9))
        self.ent_custom_ignore.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(4, 10))

        # Status bar
        sb = ttk.Frame(self.root, padding=4)
        sb.pack(fill=tk.X, side=tk.BOTTOM)
        self.var_status = tk.StringVar(value="Listo.")
        ttk.Label(sb, textvariable=self.var_status, style="Status.TLabel").pack(side=tk.LEFT, padx=4)
        self.btn_export = ttk.Button(sb, text="Exportar", command=self._export, style="TButton")
        self.btn_export.pack(side=tk.LEFT, padx=15)
        self.chk_sound = ttk.Checkbutton(sb, text="Sonido", variable=self.var_sound)
        self.chk_sound.pack(side=tk.LEFT, padx=15)
        self.var_stats = tk.StringVar(value="")
        ttk.Label(sb, textvariable=self.var_stats, style="Stats.TLabel").pack(side=tk.RIGHT, padx=8)

        # Paned
        paned = ttk.PanedWindow(self.root, orient=tk.HORIZONTAL)
        paned.pack(fill=tk.BOTH, expand=True, padx=10, pady=(0, 6))

        left = ttk.Frame(paned)
        paned.add(left, weight=1)

        self.notebook = ttk.Notebook(left)
        self.notebook.pack(fill=tk.BOTH, expand=True)

        self.tab_search = ttk.Frame(self.notebook)
        self.tab_bookmarks = ttk.Frame(self.notebook)
        self.tab_space = ttk.Frame(self.notebook)

        self.notebook.add(self.tab_search, text="Búsqueda")
        self.notebook.add(self.tab_bookmarks, text="Marcadores")
        self.notebook.add(self.tab_space, text="Espacio")
        
        # Tiny header frame for the "❓" label in search tab
        left_top = ttk.Frame(self.tab_search)
        left_top.pack(fill=tk.X, side=tk.TOP, pady=(0, 2))
        
        self.lbl_help = ttk.Label(left_top, text="❓", cursor="hand2", font=("Segoe UI", 10))
        self.lbl_help.pack(side=tk.LEFT, padx=5)

        cols = ("nombre","ruta","tipo")
        
        # Búsqueda tree
        self.tree = ttk.Treeview(self.tab_search, columns=cols, show="headings")
        for c, w, t, stretch in (("nombre", 150, "Nombre", False),
                                 ("ruta", 250, "Ruta", True),
                                 ("tipo", 70, "Tipo", False)):
            self.tree.heading(c, text=t)
            self.tree.column(c, width=w, anchor=tk.W, stretch=stretch)
        self.tree.tag_configure("noresult", foreground="red")
        sb_tree = ttk.Scrollbar(self.tab_search, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=sb_tree.set)
        sb_tree.pack(side=tk.RIGHT, fill=tk.Y)
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Marcadores tree
        self.tree_bookmarks = ttk.Treeview(self.tab_bookmarks, columns=cols, show="headings")
        for c, w, t, stretch in (("nombre", 150, "Nombre", False),
                                 ("ruta", 250, "Ruta", True),
                                 ("tipo", 70, "Tipo", False)):
            self.tree_bookmarks.heading(c, text=t)
            self.tree_bookmarks.column(c, width=w, anchor=tk.W, stretch=stretch)
        self.tree_bookmarks.tag_configure("noresult", foreground="red")
        sb_tree_bookmarks = ttk.Scrollbar(self.tab_bookmarks, orient=tk.VERTICAL, command=self.tree_bookmarks.yview)
        self.tree_bookmarks.configure(yscrollcommand=sb_tree_bookmarks.set)
        sb_tree_bookmarks.pack(side=tk.RIGHT, fill=tk.Y)
        self.tree_bookmarks.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # --- Pestaña Espacio ---
        space_top = ttk.Frame(self.tab_space, padding=5)
        space_top.pack(fill=tk.X, side=tk.TOP)
        self.btn_analyze_space = ttk.Button(space_top, text="Analizar Espacio", command=self._start_space_analysis)
        self.btn_analyze_space.pack(side=tk.LEFT, padx=5)
        self.lbl_space_status = ttk.Label(space_top, text="", font=("Segoe UI", 9, "italic"))
        self.lbl_space_status.pack(side=tk.LEFT, padx=10, fill=tk.X, expand=True)

        space_paned = ttk.PanedWindow(self.tab_space, orient=tk.VERTICAL)
        space_paned.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        self.frame_top50 = ttk.LabelFrame(space_paned, text="Top 50 Archivos más Pesados", padding=5)
        space_paned.add(self.frame_top50, weight=1)

        cols_space = ("nombre", "ruta", "tamano")
        self.tree_space = ttk.Treeview(self.frame_top50, columns=cols_space, show="headings", height=8)
        for c, w, t, stretch in (("nombre", 150, "Nombre", False),
                                 ("ruta", 250, "Ruta", True),
                                 ("tamano", 80, "Tamaño", False)):
            self.tree_space.heading(c, text=t)
            self.tree_space.column(c, width=w, anchor=tk.W, stretch=stretch)
        self.tree_space.tag_configure("noresult", foreground="red")

        sb_tree_space = ttk.Scrollbar(self.frame_top50, orient=tk.VERTICAL, command=self.tree_space.yview)
        self.tree_space.configure(yscrollcommand=sb_tree_space.set)
        sb_tree_space.pack(side=tk.RIGHT, fill=tk.Y)
        self.tree_space.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        self.frame_categories = ttk.LabelFrame(space_paned, text="Desglose por Categorías", padding=5)
        space_paned.add(self.frame_categories, weight=1)

        self.canvas_space = tk.Canvas(self.frame_categories, height=120, highlightthickness=0)
        self.canvas_space.pack(fill=tk.BOTH, expand=True)
        self.canvas_space.bind("<Configure>", lambda e: self._draw_category_chart())

        self.tree_space.bind("<<TreeviewSelect>>", self._on_select)
        self.tree_space.bind("<Double-1>", self._on_double_click)
        self.tree_space.bind("<Button-3>", self._ctx_menu)

        # Initialize Hovertips for options and help icon
        t = TRANSLATIONS[self.language]
        self.tip_help = Hovertip(self.lbl_help, t['tip_tree_help'], hover_delay=500)
        self.tip_rb_type_name = Hovertip(self.rb_type_name, t['tip_rb_type_name'], hover_delay=500)
        self.tip_rb_type_content = Hovertip(self.rb_type_content, t['tip_rb_type_content'], hover_delay=500)
        self.tip_chk_exact = Hovertip(self.chk_exact, t['tip_chk_exact'], hover_delay=500)
        self.tip_chk_whole = Hovertip(self.chk_whole, t['tip_chk_whole'], hover_delay=500)
        self.tip_chk_git = Hovertip(self.chk_ignore_git, t['tip_chk_git'], hover_delay=500)
        self.tip_chk_node = Hovertip(self.chk_ignore_node, t['tip_chk_node'], hover_delay=500)
        self.tip_chk_appdata = Hovertip(self.chk_ignore_appdata, t['tip_chk_appdata'], hover_delay=500)
        self.tip_chk_cache = Hovertip(self.chk_ignore_cache, t['tip_chk_cache'], hover_delay=500)
        self.tip_chk_system = Hovertip(self.chk_ignore_system, t['tip_chk_system'], hover_delay=500)
        self.tip_ignore_custom = Hovertip(self.ent_custom_ignore, t['tip_ignore_custom'], hover_delay=500)
        self.tip_rb_mode_all = Hovertip(self.rb_mode_all, t['tip_rb_mode_all'], hover_delay=500)
        self.tip_rb_mode_any = Hovertip(self.rb_mode_any, t['tip_rb_mode_any'], hover_delay=500)
        self.tip_cb_file_type = Hovertip(self.cb_file_type, t['tip_cb_file_type'], hover_delay=500)
        self.tip_cb_date = Hovertip(self.cb_date, t['tip_cb_date'], hover_delay=500)
        self.tip_chk_regex = Hovertip(self.chk_regex, t['tip_chk_regex'], hover_delay=500)
        self.tip_regex_help = Hovertip(self.lbl_regex_help, t['tip_regex_help'], hover_delay=500)
        self.tip_export = Hovertip(self.btn_export, t['tip_export'], hover_delay=500)
        self.tip_chk_sound = Hovertip(self.chk_sound, t['tip_chk_sound'], hover_delay=500)

        self.frame_preview = ttk.LabelFrame(paned, text="Vista Previa", padding=4)
        paned.add(self.frame_preview, weight=1)
        nav = ttk.Frame(self.frame_preview)
        nav.pack(fill=tk.X, pady=(0,4))
        self.btn_prev = ttk.Button(nav, text="< Anterior", command=self._prev, state=tk.DISABLED)
        self.btn_prev.pack(side=tk.LEFT, padx=4)
        self.lbl_nav = ttk.Label(nav, text="0 / 0", font=("Segoe UI",9,"bold"))
        self.lbl_nav.pack(side=tk.LEFT, expand=True)
        self.btn_next = ttk.Button(nav, text="Siguiente >", command=self._next, state=tk.DISABLED)
        self.btn_next.pack(side=tk.RIGHT, padx=4)
        self.txt = tk.Text(self.frame_preview, wrap=tk.WORD, state=tk.DISABLED, font=("Consolas",10), width=40)
        sb_txt = ttk.Scrollbar(self.frame_preview, orient=tk.VERTICAL, command=self.txt.yview)
        self.txt.configure(yscrollcommand=sb_txt.set)
        sb_txt.pack(side=tk.RIGHT, fill=tk.Y)
        self.txt.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.txt.tag_config("match",  background="#a8e6cf", foreground="black")
        self.txt.tag_config("active", background="#22c55e", foreground="white",
                            font=("Consolas",10,"bold"))

        self.tree.bind("<<TreeviewSelect>>", self._on_select)
        self.tree.bind("<Double-1>", self._on_double_click)
        self.tree.bind("<Button-3>", self._ctx_menu)

        self.tree_bookmarks.bind("<<TreeviewSelect>>", self._on_select)
        self.tree_bookmarks.bind("<Double-1>", self._on_double_click)
        self.tree_bookmarks.bind("<Button-3>", self._ctx_menu)

        self.cm = tk.Menu(self.root, tearoff=0)

        self._apply_theme()
        self._update_language()

    # ── helpers UI ────────────────────────────────────────────────────────────
    def _browse(self):
        d = filedialog.askdirectory()
        if d: self.var_path.set(d)

    def _on_folder_drop(self, path):
        if not path:
            return
        if os.path.isfile(path):
            folder = os.path.dirname(path)
        else:
            folder = path
        if os.path.isdir(folder):
            self.var_path.set(folder)
            t = TRANSLATIONS[self.language]
            self.var_status.set(t['status_folder_dropped'].format(path=os.path.basename(folder)))

    def _on_paste(self, event=None):
        files = get_clipboard_files()
        if files:
            self._on_folder_drop(files[0])
            return "break"
        return None

    def _get_sel_path(self):
        try:
            active_tab = self.notebook.index(self.notebook.select())
        except Exception:
            active_tab = 0
        if active_tab == 0:
            tree = self.tree
        elif active_tab == 1:
            tree = self.tree_bookmarks
        else:
            tree = self.tree_space
        sel = tree.selection()
        if not sel: return None
        v = tree.item(sel[0])['values']
        return v[1] if v and v[1] else None

    # ── animación en status bar ───────────────────────────────────────────────
    def _animate(self):
        if not self.is_searching:
            return
        t = TRANSLATIONS[self.language]
        elapsed = time.time() - self.start_time
        sc    = self.stats['scanned']
        tot   = self.stats['total']
        mat   = self.stats['matches']
        phase = self.stats.get('phase', 'recolectando')

        if phase == 'recolectando':
            tot_str = t['stats_detecting'].format(total=tot)
        else:
            tot_str = f"/ {tot}" if tot > 0 else ''

        self.var_stats.set(t['stats_running'].format(elapsed=elapsed, scanned=sc, total_str=tot_str, matches=mat))
        self.root.after(200, self._animate)

    def _poll_queue(self):
        try:
            while True:
                kind, data = self.q.get_nowait()
                if not self.is_searching and not self.is_analyzing_space:
                    if kind not in ("preview_raw", "preview_hl", "preview_image"):
                        continue
                if kind == "match":
                    score, path = data
                    if self.matches_shown < 10000:
                        t = TRANSLATIONS[self.language]
                        display_name = self._get_display_name(path)
                        tags = ()
                        bookmark = self.bookmarks.get(path)
                        if bookmark and bookmark.get("color"):
                            tags = (f"tag_{bookmark['color']}",)
                        self.tree.insert("", tk.END,
                            values=(display_name, path,
                                    t['type_folder'] if os.path.isdir(path) else t['type_file']),
                            tags=tags)
                        self.matches_shown += 1
                elif kind == "scan":
                    path, sc, tot, mc = data
                    self.stats.update({
                        'last_path': path,
                        'scanned':   sc,
                        'total':     tot,
                        'matches':   mc,
                        'phase': 'recolectando' if tot < 0 else 'procesando'
                    })
                elif kind == "status":
                    self.var_status.set(data)
                elif kind == "done":
                    self._finish(data)
                elif kind == "space_progress":
                    scanned, total_size = data
                    self.space_stats = {'scanned': scanned, 'total_size': total_size}
                    t = TRANSLATIONS[self.language]
                    self.lbl_space_status.config(text=t['status_analyzing'].format(count=scanned))
                elif kind == "space_done":
                    top50, category_sizes, elapsed, completed = data
                    self._finish_space_analysis(top50, category_sizes, elapsed, completed)
                elif kind == "preview_raw":
                    self._show_raw(data)
                elif kind == "preview_hl":
                    self._show_highlight(data)
                elif kind == "preview_image":
                    self._show_image(data)
        except queue.Empty:
            pass
        self.root.after(80, self._poll_queue)

    def _start(self):
        if self.is_searching:
            self.engine.stop()
            # Clear queue to discard any pending messages
            while not self.q.empty():
                try:
                    self.q.get_nowait()
                except queue.Empty:
                    break
            self.is_searching = False
            t = TRANSLATIONS[self.language]
            self.btn_search.config(text=t['btn_search'], state=tk.NORMAL)
            self.var_status.set(t['status_canceled'])
            self.var_stats.set(t['stats_done'].format(
                total=self.stats.get('total', 0), 
                scanned=self.stats.get('scanned', 0), 
                matches=self.stats.get('matches', 0), 
                time=time.time() - self.start_time
            ))
            return

        base = self.var_path.get().strip()
        qry  = self.var_q.get().strip()
        t = TRANSLATIONS[self.language]
        if not base or not os.path.isdir(base):
            messagebox.showwarning("Zeta", t['msg_valid_folder'])
            return
        if not qry:
            messagebox.showwarning("Zeta", t['msg_valid_query'])
            return

        # Guardar en el historial
        if qry in self.search_history:
            self.search_history.remove(qry)
        self.search_history.insert(0, qry)
        self.search_history = self.search_history[:10]
        self.cb_q.config(values=self.search_history)

        self.last_query = qry
        self.last_type  = self.var_type.get()
        for i in self.tree.get_children(): self.tree.delete(i)
        self._show_raw("")
        self.is_searching = True
        self.start_time = time.time()
        self.stats = {'scanned': 0, 'matches': 0, 'total': 0, 'last_path': '', 'phase': 'recolectando', 'elapsed_time': 0}
        self.matches_shown = 0
        self.btn_search.config(text=t['btn_cancel'])
        self._animate()

        regex = self.var_regex.get()
        if regex:
            try:
                re.compile(qry)
            except re.error as ex:
                messagebox.showwarning("Zeta", f"Expresión regular inválida: {ex}")
                self.is_searching = False
                self.btn_search.config(text=t['btn_search'])
                return
            kws = [qry]
        else:
            kws = [k.strip() for k in qry.split() if k.strip()]

        mode_and = self.var_mode.get() == "and"
        exact    = self.var_exact.get()
        whole    = self.var_whole.get()
        typ      = self.var_type.get()

        exts = FILE_TYPES[self.current_file_type_key]
        blacklist = []
        if self.var_ignore_git.get():
            blacklist.append(".git")
        if self.var_ignore_node.get():
            blacklist.append("node_modules")
        if self.var_ignore_appdata.get():
            blacklist.append("appdata")
        if self.var_ignore_cache.get():
            blacklist.append("cache")
        if self.var_ignore_system.get():
            blacklist.extend(["$recycle.bin", "system volume information", ".trash"])
        custom = self.var_custom_ignore.get().strip()
        if custom:
            for item in custom.split(","):
                val = item.strip().lower()
                if val:
                    blacklist.append(val)

        # Calcular threshold de fecha
        threshold = None
        if self.current_date_filter_key == 'date_today':
            threshold = time.time() - 24 * 3600
        elif self.current_date_filter_key == 'date_week':
            threshold = time.time() - 7 * 24 * 3600
        elif self.current_date_filter_key == 'date_month':
            threshold = time.time() - 30 * 24 * 3600

        self._save_settings()
        threading.Thread(target=self._run, args=(base, kws, typ, exact, whole, mode_and, exts, tuple(blacklist), threshold, regex),
                         daemon=True).start()

    def _run(self, base, kws, typ, exact, whole, mode_and, exts, blacklist, threshold, regex):
        t = TRANSLATIONS[self.language]
        self.q.put(("status", t['status_searching']))
        try:
            def match_cb(m):  self.q.put(("match", m))
            def scan_cb(p, sc, tot, mc): self.q.put(("scan", (p, sc, tot, mc)))

            if typ == "nombre":
                self.engine.search_name(base, kws, exact, exts, match_cb, scan_cb, blacklist, threshold, regex)
            else:
                self.engine.search_content(base, kws, mode_and, whole, exts, match_cb, scan_cb, blacklist, threshold, regex, exact)
            self.q.put(("done", not self.engine.cancel))
        except Exception as e:
            self.q.put(("status", f"Error: {e}"))
            self.q.put(("done", False))

    def _finish(self, completed):
        self.is_searching = False
        t = TRANSLATIONS[self.language]
        self.btn_search.config(text=t['btn_search'], state=tk.NORMAL)
        elapsed = time.time() - self.start_time
        self.stats['elapsed_time'] = elapsed
        m   = self.stats['matches']
        tot = self.stats['total']
        sc  = self.stats['scanned']
        if not completed:
            self.var_status.set(t['status_canceled'])
        elif m == 0:
            self.tree.insert("", tk.END, values=(t['tree_no_results'],"",""), tags=("noresult",))
            self.var_status.set(t['status_no_results'])
        else:
            self.var_status.set(t['status_results'].format(count=m, time=elapsed))
        self.var_stats.set(t['stats_done'].format(total=tot, scanned=sc, matches=m, time=elapsed))

        if completed and self.var_sound.get() and elapsed > 3.0:
            try:
                import winsound
                winsound.MessageBeep(winsound.MB_ICONASTERISK)
            except Exception:
                pass

    # ── analizador de espacio en disco ───────────────────────────────────────
    def _start_space_analysis(self):
        if self.is_analyzing_space:
            self.engine.stop()
            self.is_analyzing_space = False
            t = TRANSLATIONS[self.language]
            self.btn_analyze_space.config(text=t['btn_analyze_space'], state=tk.NORMAL)
            self.lbl_space_status.config(text=t['status_canceled'])
            return

        if self.is_searching:
            t = TRANSLATIONS[self.language]
            msg = "Por favor, cancela la búsqueda activa antes de iniciar el análisis de espacio." if self.language == 'es' else "Please cancel the active search before starting space analysis."
            messagebox.showwarning("Zeta", msg)
            return

        base = self.var_path.get().strip()
        t = TRANSLATIONS[self.language]
        if not base or not os.path.isdir(base):
            messagebox.showwarning("Zeta", t['msg_valid_folder'])
            return

        for i in self.tree_space.get_children():
            self.tree_space.delete(i)

        self.canvas_space.delete("all")

        self.is_analyzing_space = True
        self.start_time = time.time()
        self.space_stats = {'scanned': 0, 'total_size': 0}
        self.top50_data = []
        self.category_sizes = {
            'documents': 0, 'images': 0, 'media': 0, 'system': 0, 'archive': 0, 'others': 0
        }

        self.btn_analyze_space.config(text=t['btn_cancel_space'])
        self.lbl_space_status.config(text=t['status_analyzing'].format(count=0))

        blacklist = []
        if self.var_ignore_git.get():
            blacklist.append(".git")
        if self.var_ignore_node.get():
            blacklist.append("node_modules")
        if self.var_ignore_appdata.get():
            blacklist.append("appdata")
        if self.var_ignore_cache.get():
            blacklist.append("cache")
        if self.var_ignore_system.get():
            blacklist.extend(["$recycle.bin", "system volume information", ".trash"])
        custom = self.var_custom_ignore.get().strip()
        if custom:
            for item in custom.split(","):
                val = item.strip().lower()
                if val:
                    blacklist.append(val)

        threading.Thread(target=self._run_space_analysis, args=(base, tuple(blacklist)), daemon=True).start()

    def _run_space_analysis(self, base, blacklist):
        import heapq
        self.engine.cancel = False
        top_files = []
        category_sizes = {
            'documents': 0, 'images': 0, 'media': 0, 'system': 0, 'archive': 0, 'others': 0
        }
        total_size = 0
        scanned_count = 0

        def get_cat_key(ext):
            ext = ext.lower()
            if ext in FILE_TYPES['type_text']: return 'documents'
            if ext in FILE_TYPES['type_images']: return 'images'
            if ext in FILE_TYPES['type_media']: return 'media'
            if ext in FILE_TYPES['type_sys']: return 'system'
            if ext in FILE_TYPES['type_archive']: return 'archive'
            return 'others'

        try:
            for name, path, file_size in self.engine.walk_space(base, blacklist):
                if self.engine.cancel:
                    break
                
                scanned_count += 1
                total_size += file_size
                
                ext = os.path.splitext(name)[1].lower()
                cat = get_cat_key(ext)
                category_sizes[cat] += file_size
                
                if len(top_files) < 50:
                    heapq.heappush(top_files, (file_size, name, path))
                else:
                    if file_size > top_files[0][0]:
                        heapq.heapreplace(top_files, (file_size, name, path))
                        
                if scanned_count % 2000 == 0:
                    self.q.put(("space_progress", (scanned_count, total_size)))

            elapsed = time.time() - self.start_time
            sorted_top = sorted(top_files, key=lambda x: x[0], reverse=True)
            self.q.put(("space_done", (sorted_top, category_sizes, elapsed, not self.engine.cancel)))
        except Exception as e:
            self.q.put(("status", f"Error: {e}"))
            self.q.put(("space_done", ([], category_sizes, 0, False)))

    def _finish_space_analysis(self, top50, category_sizes, elapsed, completed):
        self.is_analyzing_space = False
        t = TRANSLATIONS[self.language]
        self.btn_analyze_space.config(text=t['btn_analyze_space'], state=tk.NORMAL)
        
        self.top50_data = top50
        self.category_sizes = category_sizes
        
        total_space = sum(category_sizes.values())
        self.space_stats['total_size'] = total_space
        
        if not completed:
            self.lbl_space_status.config(text=t['status_canceled'])
        else:
            self.lbl_space_status.config(text=t['status_analysis_done'].format(time=elapsed))
            
        for i in self.tree_space.get_children():
            self.tree_space.delete(i)
            
        if not top50:
            self.tree_space.insert("", tk.END, values=(t['tree_no_results'], "", ""), tags=("noresult",))
        else:
            for size, name, path in top50:
                display_name = self._get_display_name(path)
                tags = ()
                bookmark = self.bookmarks.get(path)
                if bookmark and bookmark.get("color"):
                    tags = (f"tag_{bookmark['color']}",)
                self.tree_space.insert("", tk.END, values=(display_name, path, _format_size(size)), tags=tags)
                
        self._draw_category_chart()
        
        if completed and self.var_sound.get() and elapsed > 3.0:
            try:
                import winsound
                winsound.MessageBeep(winsound.MB_ICONASTERISK)
            except Exception:
                pass

    def _draw_category_chart(self):
        if not hasattr(self, 'canvas_space'):
            return
        self.canvas_space.delete("all")
        w = self.canvas_space.winfo_width()
        if w < 200:
            w = 300
        
        t = TRANSLATIONS[self.language]
        is_dark = self.theme == 'dark'
        text_color = '#f9fafb' if is_dark else '#1f2937'
        track_color = '#374151' if is_dark else '#e5e7eb'
        
        bg_color = '#1f2937' if is_dark else '#f3f4f6'
        self.canvas_space.config(bg=bg_color)

        total_space = sum(self.category_sizes.values())
        categories = ['documents', 'images', 'media', 'system', 'archive', 'others']
        
        def draw_rounded_rect(x1, y1, x2, y2, r, fill_color):
            if x2 - x1 <= 0:
                return
            if x2 - x1 < 2 * r:
                r = (x2 - x1) / 2
            points = [
                x1 + r, y1,
                x2 - r, y1,
                x2, y1 + r,
                x2, y2 - r,
                x2 - r, y2,
                x1 + r, y2,
                x1, y2 - r,
                x1, y1 + r
            ]
            self.canvas_space.create_polygon(points, fill=fill_color, outline="", smooth=True)

        for idx, cat in enumerate(categories):
            size = self.category_sizes.get(cat, 0)
            pct = size / total_space if total_space > 0 else 0
            
            # 2-column layout: col 0 for first 3, col 1 for remaining 3
            col = 0 if idx < 3 else 1
            row = idx if idx < 3 else idx - 3
            row_y = row * 36 + 6
            
            col_width = (w - 30) / 2
            if col == 0:
                x1 = 10
                x2 = 10 + col_width
            else:
                x1 = w/2 + 5
                x2 = w - 10
            
            cat_label = t.get('cat_' + cat, cat.capitalize())
            self.canvas_space.create_text(x1, row_y + 4, text=cat_label, fill=text_color, font=("Segoe UI", 9, "bold"), anchor=tk.W)
            
            pct_str = f"{pct*100:.1f}%"
            size_str = _format_size(size)
            lbl_text = f"{size_str} ({pct_str})"
            self.canvas_space.create_text(x2, row_y + 4, text=lbl_text, fill=text_color, font=("Segoe UI", 9), anchor=tk.E)
            
            y1 = row_y + 13
            y2 = row_y + 21
            draw_rounded_rect(x1, y1, x2, y2, 4, track_color)
            
            fill_x2 = x1 + (x2 - x1) * pct
            if pct > 0:
                fill_color = self.category_colors.get(cat, '#6b7280')
                draw_rounded_rect(x1, y1, fill_x2, y2, 4, fill_color)


    # ── preview ───────────────────────────────────────────────────────────────
    def _reset_nav(self):
        self.match_indices = []
        self.cur_match = 0
        self.lbl_nav.config(text="0 / 0")
        self.btn_prev.config(state=tk.DISABLED)
        self.btn_next.config(state=tk.DISABLED)

    def _show_raw(self, text):
        self._reset_nav()
        self.txt.config(state=tk.NORMAL)
        self.txt.delete("1.0", tk.END)
        self.txt.insert(tk.END, text)
        self.txt.config(state=tk.DISABLED)

    def _show_highlight(self, data):
        self._reset_nav()
        self.txt.config(state=tk.NORMAL)
        self.txt.delete("1.0", tk.END)
        if data.get('prefix'): self.txt.insert(tk.END, data['prefix'])
        self.txt.insert(tk.END, data['content'])

        for kw in data['keywords']:
            pat = make_regex_pattern(kw)
            pos = "1.0"
            while True:
                cv = tk.StringVar()
                pos = self.txt.search(pat, pos, stopindex=tk.END, regexp=True, nocase=True, count=cv)
                if not pos: break
                ln = cv.get()
                if not ln: break
                end = f"{pos}+{ln}c"
                self.txt.tag_add("match", pos, end)
                self.match_indices.append((pos, end))
                pos = end

        self.txt.config(state=tk.DISABLED)
        self.match_indices.sort(key=lambda x: [int(n) for n in x[0].split('.')])

        if self.match_indices:
            self.cur_match = 0
            self._update_nav()
            self._hl_current()

    def _update_nav(self):
        n = len(self.match_indices)
        self.lbl_nav.config(text=f"{self.cur_match+1} / {n}")
        st = tk.NORMAL if n > 1 else tk.DISABLED
        self.btn_prev.config(state=st)
        self.btn_next.config(state=st)

    def _hl_current(self):
        if not self.match_indices: return
        self.txt.tag_remove("active","1.0",tk.END)
        s, e = self.match_indices[self.cur_match]
        self.txt.tag_add("active", s, e)
        self.txt.see(s)

    def _prev(self):
        if not self.match_indices: return
        self.cur_match = (self.cur_match - 1) % len(self.match_indices)
        self._update_nav(); self._hl_current()

    def _next(self):
        if not self.match_indices: return
        self.cur_match = (self.cur_match + 1) % len(self.match_indices)
        self._update_nav(); self._hl_current()

    def _on_select(self, _):
        path = self._get_sel_path()
        if not path or not os.path.exists(path):
            self._show_raw(""); return
        t = TRANSLATIONS[self.language]
        self._show_raw(t['preview_loading'])
        threading.Thread(target=self._load_preview, args=(path,), daemon=True).start()

    def _load_preview(self, path):
        t = TRANSLATIONS[self.language]
        try:
            if os.path.isdir(path):
                items = os.listdir(path)
                txt = t['preview_folder'].format(name=os.path.basename(path), count=len(items))
                for it in items[:40]: txt += f"  {it}\n"
                if len(items) > 40: txt += t['preview_more'].format(count=len(items)-40)
                self.q.put(("preview_raw", txt)); return

            ext = os.path.splitext(path)[1].lower()
            if ext in ('.png', '.jpg', '.jpeg', '.gif', '.bmp'):
                self.q.put(("preview_image", path))
                return

            if ext in ('.exe', '.dll', '.sys', '.msi', '.bat', '.cmd'):
                self._load_exe_sys_preview(path)
                return

            if ext in ('.pyc', '.mp3', '.zip', '.rar', '.7z', '.db', '.sqlite', '.lnk', '.bin', '.dat', '.tar', '.gz', '.bz2', '.o', '.obj'):
                self._load_other_binary_preview(path)
                return

            # Check for binary content by reading head
            try:
                with open(path, 'rb') as f:
                    head = f.read(1024)
                if b'\x00' in head:
                    self._load_other_binary_preview(path)
                    return
            except Exception:
                pass
            kws = [k.strip() for k in self.last_query.split() if k.strip()] if self.last_query else []
            if ext == '.pdf' and pypdf is not None:
                content = _extract_pdf_text(path)
                suffix = t['preview_truncated'] if len(content) > 50000 else ""
                content = content[:50000]
                if kws:
                    self.q.put(("preview_hl", {"content": content + suffix, "keywords": kws, "prefix": ""}))
                else:
                    self.q.put(("preview_raw", content + suffix))
                return
            elif ext == '.docx' and docx is not None:
                content = _extract_docx_text(path)
                suffix = t['preview_truncated'] if len(content) > 50000 else ""
                content = content[:50000]
                if kws:
                    self.q.put(("preview_hl", {"content": content + suffix, "keywords": kws, "prefix": ""}))
                else:
                    self.q.put(("preview_raw", content + suffix))
                return

            size = os.path.getsize(path)
            kws  = [k.strip() for k in self.last_query.split() if k.strip()] if self.last_query else []

            # Encontrar offset del primer match
            offset = 0
            if kws:
                kw_norm = unicodedata.normalize('NFKD', kws[0]).encode('ASCII','ignore').decode().lower()
                try:
                    with open(path,'rb') as f:
                        raw = f.read(MAX_SIZE)
                    txt_norm = unicodedata.normalize('NFKD', raw.decode('utf-8','ignore')
                                                     ).encode('ASCII','ignore').decode().lower()
                    idx = txt_norm.find(kw_norm)
                    if idx > 0:
                        # Aproximar byte offset (relación ~1:1 para ASCII)
                        offset = max(0, idx - 500)
                except Exception:
                    pass

            start = max(0, offset - 8000)
            with open(path,'rb') as f:
                f.seek(start)
                content = f.read(50000).decode('utf-8', errors='ignore')

            prefix = "…\n" if start > 0 else ""
            suffix = t['preview_truncated'] if start + 50000 < size else ""

            if kws:
                self.q.put(("preview_hl", {"content": content+suffix, "keywords": kws, "prefix": prefix}))
            else:
                self.q.put(("preview_raw", prefix + content + suffix))
        except Exception as ex:
            self.q.put(("preview_raw", t['preview_error'].format(error=ex)))

    # ── context menu / open ───────────────────────────────────────────────────
    def _ctx_menu(self, e):
        try:
            active_tab = self.notebook.index(self.notebook.select())
        except Exception:
            active_tab = 0
        tree = self.tree if active_tab == 0 else self.tree_bookmarks
        row = tree.identify_row(e.y)
        if row:
            tree.selection_set(row)
            path = self._get_sel_path()
            if path:
                self._rebuild_ctx_menu(path)
                self.cm.tk_popup(e.x_root, e.y_root)

    def _on_double_click(self, event):
        try:
            active_tab = self.notebook.index(self.notebook.select())
        except Exception:
            active_tab = 0
        tree = self.tree if active_tab == 0 else self.tree_bookmarks
        region = tree.identify_region(event.x, event.y)
        if region != "cell" and region != "tree":
            return
        row = tree.identify_row(event.y)
        if not row:
            return
        item = tree.item(row)
        if item and "noresult" in item.get("tags", ()):
            return
        self._open()

    def _get_display_name(self, path):
        basename = os.path.basename(path)
        bookmark = self.bookmarks.get(path)
        if bookmark:
            prefix = ""
            if bookmark.get("star"):
                prefix += "⭐ "
            color = bookmark.get("color")
            if color:
                color_emojis = {
                    "red": "🔴 ", "orange": "🟠 ", "yellow": "🟡 ", 
                    "green": "🟢 ", "blue": "🔵 ", "purple": "🟣 ", "gray": "⚫ "
                }
                prefix += color_emojis.get(color, "")
            return prefix + basename
        return basename

    def _refresh_bookmarks_tree(self):
        for item in self.tree_bookmarks.get_children():
            self.tree_bookmarks.delete(item)
        t = TRANSLATIONS[self.language]
        if not self.bookmarks:
            self.tree_bookmarks.insert("", tk.END, values=(t['tree_no_bookmarks'], "", ""), tags=("noresult",))
            return
        for path, bookmark in self.bookmarks.items():
            if not os.path.exists(path):
                continue
            basename = os.path.basename(path)
            is_dir = os.path.isdir(path)
            display_name = self._get_display_name(path)
            typ_str = t['type_folder'] if is_dir else t['type_file']
            tags = ()
            color = bookmark.get("color")
            if color:
                tags = (f"tag_{color}",)
            self.tree_bookmarks.insert("", tk.END, values=(display_name, path, typ_str), tags=tags)

    def _update_tree_item(self, path):
        for item in self.tree.get_children():
            values = self.tree.item(item, "values")
            if values and len(values) >= 2 and values[1] == path:
                display_name = self._get_display_name(path)
                self.tree.item(item, values=(display_name, values[1], values[2]))
                bookmark = self.bookmarks.get(path, {})
                color = bookmark.get("color")
                tags = ()
                if color:
                    tags = (f"tag_{color}",)
                self.tree.item(item, tags=tags)

    def _toggle_star(self, path):
        if path not in self.bookmarks:
            self.bookmarks[path] = {"star": True}
        else:
            self.bookmarks[path]["star"] = not self.bookmarks[path].get("star", False)
            if not self.bookmarks[path].get("star") and not self.bookmarks[path].get("color"):
                if path in self.bookmarks:
                    del self.bookmarks[path]
        self._save_settings()
        self._update_tree_item(path)
        self._refresh_bookmarks_tree()

    def _set_color_tag(self, path, color):
        if path not in self.bookmarks:
            if color:
                self.bookmarks[path] = {"color": color}
        else:
            if color:
                self.bookmarks[path]["color"] = color
            else:
                if "color" in self.bookmarks[path]:
                    del self.bookmarks[path]["color"]
                if not self.bookmarks[path].get("star") and not self.bookmarks[path].get("color"):
                    if path in self.bookmarks:
                        del self.bookmarks[path]
        self._save_settings()
        self._update_tree_item(path)
        self._refresh_bookmarks_tree()

    def _remove_bookmark(self, path):
        if path in self.bookmarks:
            del self.bookmarks[path]
        self._save_settings()
        self._update_tree_item(path)
        self._refresh_bookmarks_tree()

    def _rebuild_ctx_menu(self, path):
        t = TRANSLATIONS[self.language]
        self.cm.delete(0, tk.END)
        self.cm.add_command(label=t['menu_open'], command=self._open)
        self.cm.add_command(label=t['menu_open_loc'], command=self._open_loc)
        self.cm.add_command(label=t['menu_copy_path'], command=self._copy_path)
        self.cm.add_separator()
        
        bookmark = self.bookmarks.get(path, {})
        has_star = bookmark.get("star", False)
        star_label = "⭐ " + (t['menu_remove_star'] if has_star else t['menu_add_star'])
        self.cm.add_command(label=star_label, command=lambda: self._toggle_star(path))
        
        color_menu = tk.Menu(self.cm, tearoff=0)
        theme_colors = {
            'light': {'bg': '#f3f4f6', 'fg': '#1f2937', 'accent': '#2563eb', 'sel_fg': '#ffffff'},
            'dark': {'bg': '#1f2937', 'fg': '#f9fafb', 'accent': '#3b82f6', 'sel_fg': '#ffffff'}
        }
        c = theme_colors[self.theme]
        color_menu.config(bg=c['bg'], fg=c['fg'], activebackground=c['accent'], activeforeground=c['sel_fg'])
        
        colors = [
            ("red", "🔴", "Rojo", "Red"),
            ("orange", "🟠", "Naranja", "Orange"),
            ("yellow", "🟡", "Amarillo", "Yellow"),
            ("green", "🟢", "Verde", "Green"),
            ("blue", "🔵", "Azul", "Blue"),
            ("purple", "🟣", "Púrpura", "Purple"),
            ("gray", "⚫", "Gris", "Gray")
        ]
        
        current_color = bookmark.get("color")
        for name, emoji, es_lbl, en_lbl in colors:
            lbl = f"{emoji} {es_lbl if self.language == 'es' else en_lbl}"
            if current_color == name:
                lbl += " ✓"
            color_menu.add_command(label=lbl, command=lambda col=name: self._set_color_tag(path, col))
            
        if current_color:
            color_menu.add_separator()
            color_menu.add_command(label=t['menu_remove_tag'], command=lambda: self._set_color_tag(path, None))
            
        self.cm.add_cascade(label=t['menu_color_tags'], menu=color_menu)
        
        if bookmark:
            self.cm.add_separator()
            self.cm.add_command(label=t['menu_remove_bookmark'], command=lambda: self._remove_bookmark(path))

    def _open(self):
        p = self._get_sel_path()
        if p and os.path.exists(p):
            os.startfile(p) if sys.platform == "win32" else subprocess.run(['open', p])

    def _open_loc(self):
        p = self._get_sel_path()
        if p and os.path.exists(p):
            subprocess.Popen(['explorer', '/select,', os.path.normpath(p)])

    def _copy_path(self):
        p = self._get_sel_path()
        if p:
            self.root.clipboard_clear()
            self.root.clipboard_append(p)
            t = TRANSLATIONS[self.language]
            self.var_status.set(t['status_copied'].format(path=p))

    def _show_image(self, path):
        self._reset_nav()
        self.txt.config(state=tk.NORMAL)
        self.txt.delete("1.0", tk.END)
        try:
            from PIL import Image, ImageTk
            img = Image.open(path)
            max_w = self.txt.winfo_width() - 20
            max_h = self.txt.winfo_height() - 20
            if max_w <= 100: max_w = 300
            if max_h <= 100: max_h = 350
            
            img.thumbnail((max_w, max_h), Image.Resampling.LANCZOS)
            self.preview_photo = ImageTk.PhotoImage(img)
            self.txt.image_create(tk.END, image=self.preview_photo)
        except Exception as e:
            t = TRANSLATIONS[self.language]
            self.txt.insert(tk.END, t['preview_error'].format(error=e))
        self.txt.config(state=tk.DISABLED)

    def _export(self):
        items = self.tree.get_children()
        if not items:
            return
        first_item = self.tree.item(items[0])
        if first_item and "noresult" in first_item.get("tags", ()):
            return

        data = []
        for item in items:
            v = self.tree.item(item)['values']
            if v and len(v) >= 3:
                data.append(v)

        t = TRANSLATIONS[self.language]
        file_path = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV (valores separados por comas)", "*.csv"), ("Texto plano", "*.txt")] if self.language == "es" else 
                      [("CSV (Comma separated values)", "*.csv"), ("Plain text", "*.txt")],
            title=t['btn_export']
        )
        if file_path:
            try:
                import csv
                is_csv = file_path.lower().endswith('.csv')
                with open(file_path, 'w', encoding='utf-8', newline='') as f:
                    if is_csv:
                        delim = ';' if self.language == 'es' else ','
                        writer = csv.writer(f, delimiter=delim)
                        writer.writerow([t['col_name'], t['col_path'], t['col_type']])
                        writer.writerows(data)
                    else:
                        for name, path, typ in data:
                            f.write(f"{name}\t{path}\t{typ}\n")
                self.var_status.set(f"Exportado: {os.path.basename(file_path)}")
            except Exception as ex:
                messagebox.showerror("Zeta", f"Error al exportar: {ex}")

    # ── theme and language helpers ────────────────────────────────────────────
    def _on_language_change(self):
        align_lang = "es" if self.var_lang.get() == "Español" else "en"
        self.language = align_lang
        self._update_language()
        self._save_settings()

    def _on_date_change(self):
        t = TRANSLATIONS[self.language]
        val = self.var_date.get()
        for key in ('date_any', 'date_today', 'date_week', 'date_month'):
            if t[key] == val:
                self.current_date_filter_key = key
                break

    def _on_file_type_change(self):
        t = TRANSLATIONS[self.language]
        val = self.var_file_type.get()
        for key in ('type_all', 'type_text', 'type_images', 'type_media', 'type_sys', 'type_archive'):
            if t[key] == val:
                self.current_file_type_key = key
                break

    def _toggle_theme(self):
        self.theme = "dark" if self.theme == "light" else "light"
        self._apply_theme()
        self._save_settings()

    def _load_settings(self):
        import json
        self.config_dir = os.path.join(os.path.expanduser('~'), '.zeta_search')
        self.config_file = os.path.join(self.config_dir, 'config.json')
        self.language = "es"
        self.theme = "light"
        self.ignore_git = True
        self.ignore_node = True
        self.ignore_appdata = True
        self.ignore_cache = True
        self.ignore_system = True
        self.custom_ignore = ""
        self.search_history = []
        self.sound = True
        self.bookmarks = {}
        try:
            if os.path.exists(self.config_file):
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    cfg = json.load(f)
                    self.language = cfg.get('language', 'es')
                    self.theme = cfg.get('theme', 'light')
                    self.ignore_git = cfg.get('ignore_git', True)
                    self.ignore_node = cfg.get('ignore_node', True)
                    self.ignore_appdata = cfg.get('ignore_appdata', True)
                    self.ignore_cache = cfg.get('ignore_cache', True)
                    self.ignore_system = cfg.get('ignore_system', True)
                    self.custom_ignore = ""
                    self.search_history = cfg.get('search_history', [])
                    self.sound = cfg.get('sound', True)
                    self.bookmarks = cfg.get('bookmarks', {})
        except Exception:
            pass

    def _save_settings(self):
        import json
        try:
            if not os.path.exists(self.config_dir):
                os.makedirs(self.config_dir, exist_ok=True)
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump({
                    'language': self.language, 
                    'theme': self.theme,
                    'ignore_git': self.var_ignore_git.get() if hasattr(self, 'var_ignore_git') else self.ignore_git,
                    'ignore_node': self.var_ignore_node.get() if hasattr(self, 'var_ignore_node') else self.ignore_node,
                    'ignore_appdata': self.var_ignore_appdata.get() if hasattr(self, 'var_ignore_appdata') else self.ignore_appdata,
                    'ignore_cache': self.var_ignore_cache.get() if hasattr(self, 'var_ignore_cache') else self.ignore_cache,
                    'ignore_system': self.var_ignore_system.get() if hasattr(self, 'var_ignore_system') else self.ignore_system,
                    'custom_ignore': "",
                    'search_history': self.search_history,
                    'sound': self.var_sound.get() if hasattr(self, 'var_sound') else self.sound,
                    'bookmarks': self.bookmarks
                }, f, ensure_ascii=False, indent=4)
        except Exception:
            pass

    def _load_theme_images(self):
        from PIL import Image, ImageOps, ImageTk
        try:
            if os.path.exists(self.theme_img_path):
                img = Image.open(self.theme_img_path).convert("RGBA")
                # Light mode: original color (usually black/dark icon)
                img_light = img.resize((20, 20), Image.Resampling.LANCZOS)
                self.theme_photo_light = ImageTk.PhotoImage(img_light)
                
                # Dark mode: inverted colors (white/light icon)
                r, g, b, a = img.split()
                rgb_img = Image.merge("RGB", (r, g, b))
                inverted_rgb = ImageOps.invert(rgb_img)
                ir, ig, ib = inverted_rgb.split()
                img_dark = Image.merge("RGBA", (ir, ig, ib, a)).resize((20, 20), Image.Resampling.LANCZOS)
                self.theme_photo_dark = ImageTk.PhotoImage(img_dark)
        except Exception:
            pass

    def _update_language(self):
        lang = "es" if self.var_lang.get() == "Español" else "en"
        self.language = lang
        t = TRANSLATIONS[lang]
        
        self.root.title(t['title'])
        self.lbl_base_folder.config(text=t['lbl_base'])
        self.btn_browse.config(text=t['btn_browse'])
        self.lbl_search_text.config(text=t['lbl_search'])
        
        if not hasattr(self, 'current_file_type_key'):
            self.current_file_type_key = 'type_all'
            
        self.lbl_file_type.config(text=t['lbl_file_type'])
        file_type_values = [
            t['type_all'],
            t['type_text'],
            t['type_images'],
            t['type_media'],
            t['type_sys'],
            t['type_archive']
        ]
        self.cb_file_type.config(values=file_type_values)
        self.var_file_type.set(t[self.current_file_type_key])

        if not hasattr(self, 'current_date_filter_key'):
            self.current_date_filter_key = 'date_any'

        self.lbl_date.config(text=t['lbl_date'])
        date_values = [
            t['date_any'],
            t['date_today'],
            t['date_week'],
            t['date_month']
        ]
        self.cb_date.config(values=date_values)
        self.var_date.set(t[self.current_date_filter_key])
        
        if self.is_searching:
            curr_text = self.btn_search.cget('text')
            if curr_text in (TRANSLATIONS['es']['btn_canceling'], TRANSLATIONS['en']['btn_canceling']):
                self.btn_search.config(text=t['btn_canceling'])
            else:
                self.btn_search.config(text=t['btn_cancel'])
        else:
            self.btn_search.config(text=t['btn_search'])
            
        self.lbl_opts.config(text=t['lbl_options'])
        self.rb_type_name.config(text=t['opt_type_name'])
        self.rb_type_content.config(text=t['opt_type_content'])
        self.chk_exact.config(text=t['chk_exact'])
        self.chk_whole.config(text=t['chk_whole'])
        self.chk_regex.config(text=t['chk_regex'])
        self.lbl_ignore.config(text=t['lbl_ignore'])
        self.chk_ignore_system.config(text=t['chk_system_trash'])
        self.lbl_ignore_custom.config(text=t['lbl_ignore_custom'])
        self.rb_mode_all.config(text=t['opt_mode_all'])
        self.rb_mode_any.config(text=t['opt_mode_any'])
        self.btn_export.config(text=t['btn_export'])
        self.chk_sound.config(text=t['chk_sound'])
        
        curr_status = self.var_status.get()
        if curr_status in (TRANSLATIONS['es']['status_ready'], TRANSLATIONS['en']['status_ready']):
            self.var_status.set(t['status_ready'])
        elif curr_status in (TRANSLATIONS['es']['status_searching'], TRANSLATIONS['en']['status_searching']):
            self.var_status.set(t['status_searching'])
        elif curr_status in (TRANSLATIONS['es']['status_canceled'], TRANSLATIONS['en']['status_canceled']):
            self.var_status.set(t['status_canceled'])
        elif curr_status in (TRANSLATIONS['es']['status_no_results'], TRANSLATIONS['en']['status_no_results']):
            self.var_status.set(t['status_no_results'])
        elif "resultado(s) en" in curr_status or "match(es) in" in curr_status:
            m = self.stats.get('matches', 0)
            elapsed = self.stats.get('elapsed_time', 0)
            self.var_status.set(t['status_results'].format(count=m, time=elapsed))
        elif curr_status.startswith("Copiado: ") or curr_status.startswith("Copied: "):
            path = curr_status.split(": ", 1)[1] if ": " in curr_status else ""
            self.var_status.set(t['status_copied'].format(path=path))
            
        if self.is_searching:
            elapsed = time.time() - self.start_time
            sc = self.stats['scanned']
            tot = self.stats['total']
            mat = self.stats['matches']
            phase = self.stats.get('phase', 'recolectando')
            if phase == 'recolectando':
                tot_str = t['stats_detecting'].format(total=tot)
            else:
                tot_str = f"/ {tot}" if tot > 0 else ''
            self.var_stats.set(t['stats_running'].format(elapsed=elapsed, scanned=sc, total_str=tot_str, matches=mat))
        else:
            m = self.stats.get('matches', 0)
            tot = self.stats.get('total', 0)
            sc = self.stats.get('scanned', 0)
            elapsed = self.stats.get('elapsed_time', 0)
            self.var_stats.set(t['stats_done'].format(total=tot, scanned=sc, matches=m, time=elapsed))
            
        self.frame_preview.config(text=t['lbl_preview'])
        self.btn_prev.config(text=t['btn_prev'])
        self.btn_next.config(text=t['btn_next'])
        
        if hasattr(self, 'notebook'):
            self.notebook.tab(0, text=t['tab_search'])
            self.notebook.tab(1, text=t['tab_bookmarks'])
            if len(self.notebook.tabs()) > 2:
                self.notebook.tab(2, text=t['tab_space'])

        self.tree.heading("nombre", text=t['col_name'])
        self.tree.heading("ruta", text=t['col_path'])
        self.tree.heading("tipo", text=t['col_type'])
        if hasattr(self, 'tree_bookmarks'):
            self.tree_bookmarks.heading("nombre", text=t['col_name'])
            self.tree_bookmarks.heading("ruta", text=t['col_path'])
            self.tree_bookmarks.heading("tipo", text=t['col_type'])
            self._refresh_bookmarks_tree()
        if hasattr(self, 'tree_space'):
            self.tree_space.heading("nombre", text=t['col_name'])
            self.tree_space.heading("ruta", text=t['col_path'])
            self.tree_space.heading("tamano", text=t['col_size'])
            self.frame_top50.config(text=t['lbl_top_heaviest'])
            self.frame_categories.config(text=t['lbl_categories_breakdown'])
            self.btn_analyze_space.config(text=t['btn_cancel_space'] if self.is_analyzing_space else t['btn_analyze_space'])
            self._draw_category_chart()

        # Update Hovertips texts dynamically
        if hasattr(self, 'tip_help'):
            self.tip_help.text = t['tip_tree_help']
            self.tip_rb_type_name.text = t['tip_rb_type_name']
            self.tip_rb_type_content.text = t['tip_rb_type_content']
            self.tip_chk_exact.text = t['tip_chk_exact']
            self.tip_chk_whole.text = t['tip_chk_whole']
            self.tip_chk_git.text = t['tip_chk_git']
            self.tip_chk_node.text = t['tip_chk_node']
            self.tip_chk_appdata.text = t['tip_chk_appdata']
            self.tip_chk_cache.text = t['tip_chk_cache']
            self.tip_chk_system.text = t['tip_chk_system']
            self.tip_ignore_custom.text = t['tip_ignore_custom']
            self.tip_rb_mode_all.text = t['tip_rb_mode_all']
            self.tip_rb_mode_any.text = t['tip_rb_mode_any']
            self.tip_cb_file_type.text = t['tip_cb_file_type']
            if hasattr(self, 'tip_cb_date'):
                self.tip_cb_date.text = t['tip_cb_date']
            if hasattr(self, 'tip_chk_regex'):
                self.chk_regex.config(text=t['chk_regex'])
                self.tip_chk_regex.text = t['tip_chk_regex']
            if hasattr(self, 'tip_regex_help'):
                self.tip_regex_help.text = t['tip_regex_help']
            if hasattr(self, 'tip_export'):
                self.tip_export.text = t['tip_export']
            if hasattr(self, 'tip_chk_sound'):
                self.tip_chk_sound.text = t['tip_chk_sound']

    def _show_regex_guide(self):
        t = TRANSLATIONS[self.language]
        top = tk.Toplevel(self.root)
        top.title(t['regex_guide_title'])
        top.geometry("550x400")
        top.minsize(450, 320)
        top.transient(self.root)
        top.grab_set()

        # Set program icon
        base = sys._MEIPASS if getattr(sys, 'frozen', False) else os.path.dirname(os.path.abspath(__file__))
        ico = os.path.join(base, "buscador.ico")
        if os.path.exists(ico):
            try:
                top.iconbitmap(ico)
            except Exception:
                pass

        # Center relative to root
        self.root.update_idletasks()
        rx = self.root.winfo_x()
        ry = self.root.winfo_y()
        rw = self.root.winfo_width()
        rh = self.root.winfo_height()

        tx = rx + (rw - 550) // 2
        ty = ry + (rh - 400) // 2
        top.geometry(f"550x400+{max(0, tx)}+{max(0, ty)}")

        # Theme colors dictionary
        theme_colors = {
            'light': {
                'bg': '#f3f4f6',
                'fg': '#1f2937',
                'field_bg': '#ffffff',
            },
            'dark': {
                'bg': '#1f2937',
                'fg': '#f9fafb',
                'field_bg': '#111827',
            }
        }
        c = theme_colors[self.theme]
        top.config(bg=c['bg'])
        set_dark_titlebar(top, self.theme == 'dark')

        # Frame for Text + Scrollbar
        frame = ttk.Frame(top, padding=10)
        frame.pack(fill=tk.BOTH, expand=True)

        txt = tk.Text(frame, wrap=tk.WORD, font=("Consolas", 10), padx=8, pady=8)
        sb = ttk.Scrollbar(frame, orient=tk.VERTICAL, command=txt.yview)
        txt.configure(yscrollcommand=sb.set)

        sb.pack(side=tk.RIGHT, fill=tk.Y)
        txt.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Style Text widget based on theme
        txt.config(
            bg=c['field_bg'],
            fg=c['fg'],
            insertbackground=c['fg'],
            selectbackground=c['bg'],
            selectforeground=c['fg']
        )

        txt.insert(tk.END, t['regex_guide_text'])
        txt.config(state=tk.DISABLED)

        # Bottom Frame for Close Button
        btn_frame = ttk.Frame(top, padding=(0, 0, 0, 10))
        btn_frame.pack(fill=tk.X)
        btn_close = ttk.Button(btn_frame, text=t['btn_close'], command=top.destroy)
        btn_close.pack(side=tk.BOTTOM)

    def _apply_theme(self):
        theme_colors = {
            'light': {
                'bg': '#f3f4f6',
                'fg': '#1f2937',
                'field_bg': '#ffffff',
                'border': '#d1d5db',
                'btn_bg': '#e5e7eb',
                'btn_active': '#d1d5db',
                'btn_disabled': '#f3f4f6',
                'fg_disabled': '#9ca3af',
                'accent': '#2563eb',
                'sel_fg': '#ffffff',
                'txt_match_bg': '#a8e6cf',
                'txt_match_fg': 'black',
                'txt_match_active_bg': '#22c55e',
                'txt_match_active_fg': 'white',
                'theme_icon': '🌙'
            },
            'dark': {
                'bg': '#1f2937',
                'fg': '#f9fafb',
                'field_bg': '#111827',
                'border': '#374151',
                'btn_bg': '#374151',
                'btn_active': '#4b5563',
                'btn_disabled': '#1f2937',
                'fg_disabled': '#6b7280',
                'accent': '#3b82f6',
                'sel_fg': '#ffffff',
                'txt_match_bg': '#1e3a8a',
                'txt_match_fg': '#f9fafb',
                'txt_match_active_bg': '#2563eb',
                'txt_match_active_fg': '#ffffff',
                'theme_icon': '☀️'
            }
        }
        
        c = theme_colors[self.theme]
        
        self.root.config(bg=c['bg'])
        set_dark_titlebar(self.root, self.theme == 'dark')
        self.txt.config(
            bg=c['field_bg'], 
            fg=c['fg'], 
            insertbackground=c['fg'], 
            selectbackground=c['accent'], 
            selectforeground=c['sel_fg']
        )
        self.txt.tag_config("match", background=c['txt_match_bg'], foreground=c['txt_match_fg'])
        self.txt.tag_config("active", background=c['txt_match_active_bg'], foreground=c['txt_match_active_fg'])
        self.cm.config(bg=c['bg'], fg=c['fg'], activebackground=c['accent'], activeforeground=c['sel_fg'])
        
        self.style.configure('.', background=c['bg'], foreground=c['fg'])
        self.style.configure('TFrame', background=c['bg'])
        self.style.configure('TLabel', background=c['bg'], foreground=c['fg'])
        self.style.configure('TLabelframe', background=c['bg'], foreground=c['fg'], bordercolor=c['border'])
        self.style.configure('TLabelframe.Label', background=c['bg'], foreground=c['fg'])
        
        self.style.configure('TButton', background=c['btn_bg'], foreground=c['fg'], bordercolor=c['border'], focuscolor='', lightcolor=c['btn_bg'], darkcolor=c['btn_bg'])
        self.style.map('TButton', 
            background=[('active', c['btn_active']), ('disabled', c['btn_disabled'])], 
            foreground=[('disabled', c['fg_disabled'])],
            bordercolor=[('active', c['border']), ('disabled', c['border'])]
        )
        
        self.style.configure('TRadiobutton', background=c['bg'], foreground=c['fg'], focuscolor='', indicatorbackground=c['field_bg'], indicatorforeground=c['fg'])
        self.style.map('TRadiobutton', 
            background=[('active', c['bg'])], 
            indicatorbackground=[('selected', c['accent']), ('!selected', c['field_bg'])],
            foreground=[('active', c['fg'])]
        )
        self.style.configure('TCheckbutton', background=c['bg'], foreground=c['fg'], focuscolor='', indicatorbackground=c['field_bg'], indicatorforeground=c['fg'])
        self.style.map('TCheckbutton', 
            background=[('active', c['bg'])], 
            indicatorbackground=[('selected', c['accent']), ('!selected', c['field_bg'])],
            foreground=[('active', c['fg'])]
        )
        
        self.style.configure('TEntry', fieldbackground=c['field_bg'], foreground=c['fg'], bordercolor=c['border'], lightcolor=c['field_bg'], darkcolor=c['field_bg'])
        self.style.map('TEntry', bordercolor=[('focus', c['accent']), ('!focus', c['border'])])
        
        self.style.configure('Treeview', background=c['field_bg'], foreground=c['fg'], fieldbackground=c['field_bg'], bordercolor=c['border'])
        self.style.configure('Treeview.Heading', background=c['btn_bg'], foreground=c['fg'], bordercolor=c['border'], lightcolor=c['btn_bg'], darkcolor=c['btn_bg'])
        self.style.map('Treeview.Heading', background=[('active', c['btn_active'])])
        self.style.map('Treeview', background=[('selected', c['accent'])], foreground=[('selected', c['sel_fg'])])
        
        self.style.configure('TCombobox', fieldbackground=c['field_bg'], foreground=c['fg'], background=c['btn_bg'], bordercolor=c['border'], arrowcolor=c['fg'])
        self.style.map('TCombobox', 
            fieldbackground=[('readonly', c['field_bg'])], 
            foreground=[('readonly', c['fg'])],
            arrowcolor=[('active', c['accent'])]
        )

        self.style.configure('TScrollbar',
            troughcolor=c['bg'],
            background=c['btn_bg'],
            bordercolor=c['border'],
            arrowcolor=c['fg'],
            lightcolor=c['bg'],
            darkcolor=c['bg']
        )
        self.style.map('TScrollbar',
            background=[('active', c['btn_active']), ('disabled', c['bg'])],
            arrowcolor=[('active', c['accent'])]
        )
        
        self.style.configure("Status.TLabel", font=("Segoe UI", 10, "bold"), background=c['bg'], foreground=c['fg'])
        self.style.configure("Stats.TLabel", font=("Segoe UI", 9), background=c['bg'], foreground=c['fg'] if self.theme == 'dark' else '#333333')
        
        # Update theme button icon or text
        if self.theme == 'light' and self.theme_photo_light:
            self.btn_theme.config(image=self.theme_photo_light, text="")
        elif self.theme == 'dark' and self.theme_photo_dark:
            self.btn_theme.config(image=self.theme_photo_dark, text="")
        else:
            self.btn_theme.config(image="", text=c['theme_icon'])
            
        self.tree.tag_configure("noresult", foreground="red" if self.theme == 'light' else "#ff6b6b")
        if hasattr(self, 'tree_bookmarks'):
            self.tree_bookmarks.tag_configure("noresult", foreground="red" if self.theme == 'light' else "#ff6b6b")
        if hasattr(self, 'tree_space'):
            self.tree_space.tag_configure("noresult", foreground="red" if self.theme == 'light' else "#ff6b6b")
            
        tag_colors = {
            'light': {
                'red': ('#fee2e2', '#991b1b'),
                'orange': ('#ffedd5', '#9a3412'),
                'yellow': ('#fef9c3', '#854d0e'),
                'green': ('#dcfce7', '#166534'),
                'blue': ('#dbeafe', '#1e40af'),
                'purple': ('#f3e8ff', '#6b21a8'),
                'gray': ('#f3f4f6', '#374151')
            },
            'dark': {
                'red': ('#7f1d1d', '#fecaca'),
                'orange': ('#7c2d12', '#ffedd5'),
                'yellow': ('#713f12', '#fef9c3'),
                'green': ('#064e3b', '#dcfce7'),
                'blue': ('#1e3a8a', '#dbeafe'),
                'purple': ('#581c87', '#f3e8ff'),
                'gray': ('#374151', '#f3f4f6')
            }
        }
        
        for col, (bg, fg) in tag_colors[self.theme].items():
            self.tree.tag_configure(f"tag_{col}", background=bg, foreground=fg)
            if hasattr(self, 'tree_bookmarks'):
                self.tree_bookmarks.tag_configure(f"tag_{col}", background=bg, foreground=fg)
            if hasattr(self, 'tree_space'):
                self.tree_space.tag_configure(f"tag_{col}", background=bg, foreground=fg)

        if hasattr(self, 'canvas_space'):
            self._draw_category_chart()


    def _get_file_metadata(self, filepath):
        metadata = {}
        if sys.platform != "win32" or version_dll is None:
            return metadata
        try:
            dw_handle = wintypes.DWORD(0)
            size = GetFileVersionInfoSizeW(filepath, ctypes.byref(dw_handle))
            if size == 0:
                return metadata
                
            buf = ctypes.create_string_buffer(size)
            if not GetFileVersionInfoW(filepath, 0, size, buf):
                return metadata
                
            p_translation = ctypes.c_void_p()
            translation_len = ctypes.c_uint(0)
            lang_codes = []
            
            if VerQueryValueW(buf, "\\VarFileInfo\\Translation", ctypes.byref(p_translation), ctypes.byref(translation_len)):
                if translation_len.value >= 4:
                    ptr = ctypes.cast(p_translation, ctypes.POINTER(wintypes.WORD))
                    for i in range(translation_len.value // 4):
                        lang_id = ptr[i*2]
                        code_page = ptr[i*2+1]
                        lang_codes.append(f"{lang_id:04x}{code_page:04x}")
            
            if not lang_codes:
                lang_codes = ["040904b0", "040a04b0", "040904e4", "04090000"]
                
            keys = [
                ("CompanyName", "Empresa creadora" if self.language == "es" else "Company"),
                ("FileDescription", "Descripción" if self.language == "es" else "Description"),
                ("FileVersion", "Versión de archivo" if self.language == "es" else "File Version"),
                ("LegalCopyright", "Copyright" if self.language == "es" else "Copyright"),
                ("ProductName", "Producto" if self.language == "es" else "Product Name"),
                ("ProductVersion", "Versión de producto" if self.language == "es" else "Product Version"),
            ]
            
            for key, label in keys:
                for lang_code in lang_codes:
                    p_value = ctypes.c_void_p()
                    value_len = ctypes.c_uint(0)
                    query_path = f"\\StringFileInfo\\{lang_code}\\{key}"
                    if VerQueryValueW(buf, query_path, ctypes.byref(p_value), ctypes.byref(value_len)):
                        if value_len.value > 0:
                            value_str = ctypes.wstring_at(p_value)
                            if value_str:
                                metadata[label] = value_str.strip()
                                break
        except Exception:
            pass
        return metadata

    def _get_signature_info(self, path):
        if sys.platform != "win32":
            return None
        try:
            cmd = [
                "powershell", "-NoProfile", "-NonInteractive", "-ExecutionPolicy", "Bypass",
                "-Command", f"Get-AuthenticodeSignature '{path}' | Select-Object Status, StatusMessage, SignerCertificate | ConvertTo-Json"
            ]
            res = subprocess.run(cmd, capture_output=True, text=True, timeout=2, creationflags=0x08000000)
            if res.returncode == 0 and res.stdout.strip():
                import json
                data = json.loads(res.stdout)
                status = data.get("Status")
                status_msg = data.get("StatusMessage")
                cert = data.get("SignerCertificate")
                subject = ""
                if cert:
                    if isinstance(cert, dict):
                        subject = cert.get("Subject", "")
                    elif isinstance(cert, str):
                        subject = cert
                        
                status_mapping = {
                    0: "Válida" if self.language == 'es' else "Valid",
                    1: "Error desconocido" if self.language == 'es' else "Unknown error",
                    2: "No firmado" if self.language == 'es' else "Not signed",
                    3: "No confiable" if self.language == 'es' else "Not trusted",
                    4: "Firma corrupta / Hash incorrecto" if self.language == 'es' else "Hash mismatch",
                    5: "Incompatible" if self.language == 'es' else "Incompatible",
                }
                
                status_text = status_mapping.get(status, str(status))
                if status == 0:
                    return {
                        "status": "Válida" if self.language == 'es' else "Valid",
                        "signer": subject or ("Desconocido" if self.language == 'es' else "Unknown")
                    }
                elif status == 2:
                    return {
                        "status": "No firmado" if self.language == 'es' else "Not signed",
                        "signer": None
                    }
                else:
                    return {
                        "status": f"{status_text} ({status_msg})" if status_msg else status_text,
                        "signer": subject or None
                    }
        except Exception:
            pass
        return None

    def _generate_hex_dump(self, filepath, max_bytes=512):
        try:
            with open(filepath, 'rb') as f:
                data = f.read(max_bytes)
            if not data:
                return ""
                
            lines = []
            lines.append("─" * 60)
            lines.append("VOLCADO HEXADECIMAL (Primeros 512 bytes)" if self.language == 'es' else "HEX DUMP (First 512 bytes)")
            lines.append("─" * 60)
            
            for i in range(0, len(data), 16):
                chunk = data[i:i+16]
                hex_str = " ".join(f"{b:02X}" for b in chunk)
                if len(chunk) < 16:
                    hex_str += " " * (3 * (16 - len(chunk)))
                ascii_str = "".join(chr(b) if 32 <= b < 127 else "." for b in chunk)
                lines.append(f"{i:08X}  {hex_str}  |{ascii_str}|")
            return "\n".join(lines)
        except Exception as e:
            return f"Error reading hex: {e}"

    def _load_exe_sys_preview(self, path):
        t = TRANSLATIONS[self.language]
        try:
            basename = os.path.basename(path)
            header = f"=== DETALLES DEL ARCHIVO: {basename} ===\n" if self.language == 'es' else f"=== FILE DETAILS: {basename} ===\n"
            
            size_bytes = os.path.getsize(path)
            if size_bytes < 1024:
                size_str = f"{size_bytes} B"
            elif size_bytes < 1024 * 1024:
                size_str = f"{size_bytes / 1024:.2f} KB"
            else:
                size_str = f"{size_bytes / (1024 * 1024):.2f} MB"
                
            info_lines = []
            info_lines.append(f"{'Tamaño' if self.language == 'es' else 'Size'}: {size_str}")
            
            metadata = self._get_file_metadata(path)
            if metadata:
                for key, val in metadata.items():
                    info_lines.append(f"{key}: {val}")
                    
            sig = self._get_signature_info(path)
            if sig:
                info_lines.append(f"{'Firma Digital' if self.language == 'es' else 'Digital Signature'}: {sig['status']}")
                if sig['signer']:
                    info_lines.append(f"{'Firmante' if self.language == 'es' else 'Signer'}: {sig['signer']}")
                    
            info_text = header + "\n".join(info_lines) + "\n\n"
            hex_dump = self._generate_hex_dump(path)
            self.q.put(("preview_raw", info_text + hex_dump))
        except Exception as ex:
            self.q.put(("preview_raw", t['preview_error'].format(error=ex)))

    def _load_other_binary_preview(self, path):
        t = TRANSLATIONS[self.language]
        try:
            basename = os.path.basename(path)
            size_bytes = os.path.getsize(path)
            if size_bytes < 1024:
                size_str = f"{size_bytes} B"
            elif size_bytes < 1024 * 1024:
                size_str = f"{size_bytes / 1024:.2f} KB"
            else:
                size_str = f"{size_bytes / (1024 * 1024):.2f} MB"
                
            header = f"=== ARCHIVO BINARIO: {basename} ===\n" if self.language == 'es' else f"=== BINARY FILE: {basename} ===\n"
            info_text = header + f"{'Tamaño' if self.language == 'es' else 'Size'}: {size_str}\n\n"
            
            hex_dump = self._generate_hex_dump(path)
            self.q.put(("preview_raw", info_text + hex_dump))
        except Exception as ex:
            self.q.put(("preview_raw", t['preview_error'].format(error=ex)))


if __name__ == '__main__':
    multiprocessing.freeze_support()   # obligatorio para PyInstaller + multiprocessing
    root = tk.Tk()
    ZetaApp(root)
    root.mainloop()
